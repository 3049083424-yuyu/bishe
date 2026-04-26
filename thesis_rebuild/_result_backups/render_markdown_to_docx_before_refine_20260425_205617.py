from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


TOKEN_RE = re.compile(r"(\*\*.*?\*\*|\[\d+(?:-\d+)?\])")
IMAGE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
REFERENCE_RE = re.compile(r"^\[\d+\]\s*")
FIXED_HEADER_TEXT = "中国石油大学（北京）本科毕业设计(论文)"
FRONT_MATTER_HEADINGS = {"摘要", "Abstract", "目录"}
FIXED_HEADER_SECTIONS = {"参考文献", "致谢", "附录"}


def set_run_font(
    run,
    size_pt: float,
    *,
    bold: bool = False,
    superscript: bool = False,
    east_asia: str = "宋体",
    latin: str = "Times New Roman",
) -> None:
    run.bold = bold
    run.font.superscript = superscript
    run.font.size = Pt(size_pt)
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_style(
    paragraph,
    *,
    first_line_indent: bool = True,
    line_spacing: float = 1.25,
    left_indent_pt: float = 0,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.left_indent = Pt(left_indent_pt)
    fmt.first_line_indent = Pt(24) if first_line_indent else Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def clear_paragraph(paragraph) -> None:
    paragraph.text = ""
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)


def configure_document(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.25
    enable_update_fields_on_open(doc)


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure_section(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(20)


def set_page_number_format(section, *, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    elif qn("w:start") in pg_num_type.attrib:
        del pg_num_type.attrib[qn("w:start")]


def add_field_run(paragraph, instruction: str, *, placeholder: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)
    return run


def configure_header_footer(section, *, header_text: str, page_fmt: str, page_start: int | None = None) -> None:
    configure_section(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    set_page_number_format(section, fmt=page_fmt, start=page_start)

    header_para = section.header.paragraphs[0]
    clear_paragraph(header_para)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(header_para, first_line_indent=False)
    header_run = header_para.add_run(header_text)
    set_run_font(header_run, 10.5)

    footer_para = section.footer.paragraphs[0]
    clear_paragraph(footer_para)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(footer_para, first_line_indent=False)
    page_run = add_field_run(footer_para, "PAGE", placeholder="1")
    set_run_font(page_run, 10.5, east_asia="Times New Roman", latin="Times New Roman")


def create_new_section(
    doc: Document,
    *,
    header_text: str,
    page_fmt: str,
    page_start: int | None = None,
):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_header_footer(section, header_text=header_text, page_fmt=page_fmt, page_start=page_start)
    return section


def add_inline_markdown(paragraph, text: str, *, size_pt: float = 12, first_line_indent: bool = True) -> None:
    set_paragraph_style(paragraph, first_line_indent=first_line_indent)
    parts = TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size_pt, bold=True)
        elif re.fullmatch(r"\[\d+(?:-\d+)?\]", part):
            run = paragraph.add_run(part)
            set_run_font(run, 10.5, superscript=True, east_asia="Times New Roman")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size_pt)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    set_run_font(run, 18, bold=True, east_asia="黑体")


def add_centered_text(doc: Document, text: str, *, size_pt: float = 14, bold: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(p, first_line_indent=False)
    run = p.add_run(text)
    set_run_font(run, size_pt, bold=bold, east_asia="黑体" if bold else "宋体")


def add_keyword_paragraph(doc: Document, label: str, content: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_style(p, first_line_indent=False)
    label_run = p.add_run(label)
    set_run_font(label_run, 12, bold=True)
    content_run = p.add_run(content)
    set_run_font(content_run, 12)


def add_heading_level_1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, 16, bold=True, east_asia="黑体")


def add_heading_level_2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, 14, bold=True, east_asia="黑体")


def add_heading_level_3(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 3")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True, east_asia="黑体")


def parse_table_cells(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    candidate = line.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
    return candidate == ""


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = [parse_table_cells(line) for line in lines if line.strip()]
    if len(rows) < 2:
        return

    header = rows[0]
    data_rows = rows[2:] if len(rows) >= 2 and is_table_separator(lines[1]) else rows[1:]
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col, cell_text in enumerate(header):
        cell = table.cell(0, col)
        p = cell.paragraphs[0]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text)
        set_run_font(run, 10.5, bold=True, east_asia="黑体")

    for row_idx, row in enumerate(data_rows, start=1):
        for col, cell_text in enumerate(row):
            cell = table.cell(row_idx, col)
            p = cell.paragraphs[0]
            clear_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            set_run_font(run, 10.5)

    doc.add_paragraph()


def resolve_image_path(markdown_path: Path, raw_path: str) -> Path:
    path = Path(raw_path.strip())
    return path if path.is_absolute() else (markdown_path.parent / path).resolve()


def add_image(doc: Document, markdown_path: Path, alt_text: str, raw_path: str) -> None:
    image_path = resolve_image_path(markdown_path, raw_path)
    if not image_path.exists():
        p = doc.add_paragraph()
        add_inline_markdown(p, f"[缺失图片] {alt_text}: {image_path}", size_pt=11, first_line_indent=False)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Mm(150))

    if alt_text.strip():
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_style(caption, first_line_indent=False)
        cap_run = caption.add_run(alt_text.strip())
        set_run_font(cap_run, 10.5)


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_style(p, first_line_indent=False)
    toc_run = add_field_run(p, r'TOC \o "1-3" \h \z \u', placeholder="目录将在打开文档后更新")
    set_run_font(toc_run, 12)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_style(note, first_line_indent=False)
    run = note.add_run("提示：如目录未自动显示，请在 Word/WPS 中更新域。")
    set_run_font(run, 10.5)


def add_reference_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = 1.25
    fmt.left_indent = Pt(24)
    fmt.first_line_indent = Pt(-24)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 12)


def is_body_chapter_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\s", text))


def render_markdown(markdown_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    configure_section(doc.sections[0])

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    front_heading_count = 0
    front_section_created = False
    body_started = False
    in_reference_block = False
    current_h2_heading = ""
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            add_image(doc, markdown_path, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if stripped.startswith("# "):
            add_title(doc, stripped[2:].strip())
            create_new_section(
                doc,
                header_text=FIXED_HEADER_TEXT,
                page_fmt="upperRoman",
                page_start=1,
            )
            front_section_created = True
            i += 1
            continue

        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            in_reference_block = False
            current_h2_heading = heading_text

            if heading_text == "目录":
                if front_heading_count > 0:
                    doc.add_page_break()
                add_heading_level_1(doc, heading_text)
                add_toc_field(doc)
                front_heading_count += 1
                i += 1
                while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("## "):
                    i += 1
                continue

            if heading_text in FRONT_MATTER_HEADINGS:
                if not front_section_created:
                    create_new_section(
                        doc,
                        header_text=FIXED_HEADER_TEXT,
                        page_fmt="upperRoman",
                        page_start=1,
                    )
                    front_section_created = True
                elif front_heading_count > 0:
                    doc.add_page_break()
                add_heading_level_1(doc, heading_text)
                front_heading_count += 1
                i += 1
                continue

            if is_body_chapter_heading(heading_text) or heading_text in FIXED_HEADER_SECTIONS:
                header_text = heading_text if heading_text not in FIXED_HEADER_SECTIONS else FIXED_HEADER_TEXT
                page_start = 1 if not body_started else None
                create_new_section(
                    doc,
                    header_text=header_text,
                    page_fmt="decimal",
                    page_start=page_start,
                )
                body_started = True
                add_heading_level_1(doc, heading_text)
                in_reference_block = heading_text == "参考文献"
                i += 1
                continue

            create_new_section(
                doc,
                header_text=FIXED_HEADER_TEXT,
                page_fmt="decimal",
                page_start=None if body_started else 1,
            )
            body_started = True
            add_heading_level_1(doc, heading_text)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading_level_2(doc, stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith("#### "):
            add_heading_level_3(doc, stripped[5:].strip())
            i += 1
            continue

        if in_reference_block and REFERENCE_RE.match(stripped):
            add_reference_paragraph(doc, stripped)
            i += 1
            continue

        if stripped.startswith("**关键词：**"):
            add_keyword_paragraph(doc, "关键词：", stripped[len("**关键词：**") :].strip())
            i += 1
            continue

        if stripped.startswith("**Key Words:**"):
            add_keyword_paragraph(doc, "Key Words: ", stripped[len("**Key Words:**") :].strip())
            i += 1
            continue

        if current_h2_heading == "Abstract" and re.fullmatch(r"\*\*.*\*\*", stripped):
            add_centered_text(doc, stripped[2:-2].strip(), size_pt=14, bold=True)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_markdown(p, stripped, size_pt=12, first_line_indent=not stripped.startswith("**"))
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--input",
        default=str(Path(r"D:/graduate/thesis_rebuild/manuscript/thesis_draft_dual_key_2011_2025.md")),
    )
    parser.add_argument(
        "--output",
        default=str(Path(r"D:/graduate/thesis_rebuild/manuscript/thesis_draft_dual_key_2011_2025.docx")),
    )
    args = parser.parse_args()
    render_markdown(Path(args.input), Path(args.output))


if __name__ == "__main__":
    main()
