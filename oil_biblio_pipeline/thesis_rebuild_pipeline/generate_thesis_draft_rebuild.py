from __future__ import annotations

import csv
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REBUILD_ROOT = Path(r"D:\graduate\thesis_rebuild")
MANUSCRIPT_DIR = REBUILD_ROOT / "manuscript"
QA_DIR = REBUILD_ROOT / "qa"
CORPUS_DIR = REBUILD_ROOT / "corpus"
INSTITUTION_DIR = REBUILD_ROOT / "institution_eval"
TOPIC_DIR = REBUILD_ROOT / "topic_evolution_dual_key_2011_2025"
OPENING_REPORT_PATH = Path(r"D:\graduate\2022012085-余冰雁-开题报告.docx")

TITLE_ZH = "基于高质量期刊文献挖掘的全球石油领域研究机构学术影响力评价与主题演化分析"
TITLE_EN = (
    "Academic Influence Evaluation and Topic Evolution Analysis of Global Petroleum "
    "Research Institutions Based on High-Quality Journal Literature Mining"
)

MD_OUTPUT = MANUSCRIPT_DIR / "thesis_draft_dual_key_2011_2025.md"
DOCX_OUTPUT = MANUSCRIPT_DIR / "thesis_draft_dual_key_2011_2025.docx"
QA_OUTPUT = QA_DIR / "thesis_delivery_qa_2026-04-24.md"

KEYWORDS_ZH = ["高质量期刊文献", "石油领域", "研究机构", "学术影响力", "主题演化"]
KEYWORDS_EN = [
    "high-quality journal literature",
    "petroleum research institutions",
    "academic influence",
    "TOPSIS",
    "topic evolution",
]

ENGLISH_NAME_MAP = {
    "中国科学院": "Chinese Academy of Sciences",
    "中国石油大学（北京）": "China University of Petroleum (Beijing)",
}

REFERENCES = [
    "Schlögl C, Stock W G, Reichmann G. Scientometric evaluation of research institutions: "
    "Identifying the appropriate dimensions and attributes for assessment[J]. Journal of "
    "Information Science Theory and Practice, 2025, 13(2): 49-68.",
    "Deng Z, Duan Z, Zhang Z, et al. An evaluation model for authors' academic influence "
    "based on multi-source heterogeneous database in bilingual environment[J]. Journal of "
    "Physics: Conference Series, 2020, 1575(1): 012147.",
    "林子婕, 唐星龙. 互动视角下学者学术影响力多维评价模型研究[J]. 情报理论与实践, "
    "2024, 47(9): 88-98.",
    "张璜, 高睿. 主题性视域下中国设计学学者学术影响力的评价立场: 基于LDA模型的实证研究[J]. "
    "上海视觉, 2024(4): 136-143.",
    "王宏宇, 石锴文, 王晓光, 等. 基于词向量网络的科研主题演化分析: 语义漂移过程的揭示[J]. "
    "情报学报, 2025, 44(10): 1287-1299.",
    "李小燕, 郑军卫, 田欣, 等. 中文科技期刊影响力分析与提升路径: 以石油天然气工程类期刊为例[J]. "
    "中国科技期刊研究, 2016, 27(11): 1221-1227.",
    "Ghosh A, Das S. Research productivity of University of Petroleum and Energy Studies "
    "during 2004-2018: A scientometric analysis[J]. Library Philosophy and Practice, 2020: 1-10.",
    "J B, T B. Mapping the research productivity in University of Petroleum and Energy Studies: "
    "A scientometric approach[J]. Library Philosophy and Practice, 2019.",
    "Hwang C L, Yoon K. Multiple Attribute Decision Making: Methods and Applications[M]. "
    "Berlin: Springer, 1981.",
    "Shannon C E. A mathematical theory of communication[J]. Bell System Technical Journal, "
    "1948, 27(3): 379-423.",
    "Diakoulaki D, Mavrotas G, Papayannakis L. Determining objective weights in multiple "
    "criteria problems: The CRITIC method[J]. Computers & Operations Research, 1995, 22(7): 763-770.",
    "Blei D M, Ng A Y, Jordan M I. Latent Dirichlet allocation[J]. Journal of Machine Learning "
    "Research, 2003, 3: 993-1022.",
    "Halevi G, Moed H, Bar-Ilan J. Suitability of Google Scholar as a source of scientific "
    "information and as a source of data for scientific evaluation: Review of the literature[J]. "
    "Journal of Informetrics, 2017, 11(3): 823-834.",
    "Gusenbauer M. Beyond Google Scholar, Scopus, and Web of Science: An evaluation of the "
    "backward and forward citation coverage of 59 databases' citation indices[J]. Research "
    "Synthesis Methods, 2024, 15(5): 802-817.",
    "Zhao R, Wang X, Liu Z, et al. Research on the impact evaluation of academic journals "
    "based on altmetrics and citation indicators[J]. Proceedings of the Association for "
    "Information Science and Technology, 2019, 56(1): 336-345.",
    "Pratama I B, Wijaya A, Hermawan B, et al. Evaluating academic performance and scholarly "
    "impact of rectors of Indonesia's public universities: A dual bibliometric and scholastic "
    "analysis[J]. Cogent Education, 2024, 11(1).",
    "Xiao-Jun H, Jian-Hong L, Ronald R. A warning for Chinese academic evaluation systems: "
    "Short-term bibliometric measures misjudge the value of pioneering contributions[J]. "
    "Journal of Zhejiang University Science B, 2018, 19(1): 1-5.",
    "Qosimjonov S A. Scientometric indicators as tools for evaluating innovation and research "
    "productivity[J]. Technical Science Integrated Research, 2025, 1(3): 24-29.",
]


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    for encoding in ("gb18030", "utf-8-sig", "utf-8"):
        try:
            with path.open("r", encoding=encoding, newline="") as handle:
                return list(csv.DictReader(handle))
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("unknown", b"", 0, 1, f"无法识别编码: {path}")


def as_float(value: str) -> float:
    return float(value.strip()) if value and value.strip() else 0.0


def as_int(value: str) -> int:
    return int(float(value.strip())) if value and value.strip() else 0


def join_cn(items: list[str], sep: str = "、") -> str:
    cleaned = [item for item in items if item]
    return sep.join(cleaned)


def md_image_path(path: Path) -> str:
    return "../" + path.relative_to(REBUILD_ROOT).as_posix()


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def load_cover_info() -> dict[str, str]:
    info = {
        "课题名称": TITLE_ZH,
        "学号": "2022012085",
        "姓名": "余冰雁",
        "学院": "经济管理学院",
        "学科专业": "信息管理与信息系统",
    }
    if not OPENING_REPORT_PATH.exists():
        return info
    try:
        report = Document(OPENING_REPORT_PATH)
        if not report.tables:
            return info
        for row in report.tables[0].rows:
            if len(row.cells) < 2:
                continue
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            if key and value:
                info[key] = value
    except Exception:
        return info
    return info


def top_n_by_group(
    rows: list[dict[str, str]],
    group_field: str,
    value_field: str,
    n: int,
) -> dict[str, list[dict[str, str]]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        grouped[row[group_field]].append(row)
    for items in grouped.values():
        items.sort(key=lambda item: as_float(item[value_field]), reverse=True)
    return {key: value[:n] for key, value in grouped.items()}


def load_data() -> dict[str, object]:
    merged_rows = read_csv_rows(CORPUS_DIR / "merged_clean_dual_key_dedup_2011_2025.csv")
    top20 = read_csv_rows(INSTITUTION_DIR / "institution_topsis_top20_dual_key_2011_2025.csv")
    weights = read_csv_rows(INSTITUTION_DIR / "institution_weight_scheme_top100_dual_key_2011_2025.csv")
    dimension_rows = read_csv_rows(INSTITUTION_DIR / "institution_dimension_top10_dual_key_2011_2025.csv")
    type_review = read_csv_rows(INSTITUTION_DIR / "institution_type_review_top100_dual_key_2011_2025.csv")

    preprocess_rows = read_csv_rows(TOPIC_DIR / "topic_preprocess_stats_dual_key_2011_2025.csv")
    model_rows = read_csv_rows(TOPIC_DIR / "topic_model_selection_dual_key_2011_2025.csv")
    strength_rows = read_csv_rows(TOPIC_DIR / "topic_strength_dual_key_2011_2025.csv")
    evolution_rows = read_csv_rows(TOPIC_DIR / "topic_evolution_paths_dual_key_2011_2025.csv")
    type_dist_rows = read_csv_rows(TOPIC_DIR / "topic_distribution_by_institution_type_dual_key_2011_2025.csv")
    level_dist_rows = read_csv_rows(TOPIC_DIR / "topic_distribution_by_institution_level_dual_key_2011_2025.csv")
    profile_rows = read_csv_rows(TOPIC_DIR / "institution_profile_classification_dual_key_2011_2025.csv")

    selected_models = {
        row["阶段"]: row
        for row in model_rows
        if row["是否选中"] == "1"
    }
    strength_top3 = top_n_by_group(strength_rows, "阶段", "主题强度", 3)

    type_grouped: dict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for row in type_dist_rows:
        type_grouped[(row["阶段"], row["机构类型"])].append(row)
    for items in type_grouped.values():
        items.sort(key=lambda item: as_float(item["主题占比"]), reverse=True)

    level_grouped: dict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for row in level_dist_rows:
        level_grouped[(row["阶段"], row["机构层级"])].append(row)
    for items in level_grouped.values():
        items.sort(key=lambda item: as_float(item["主题占比"]), reverse=True)

    dimension_grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in dimension_rows:
        dimension_grouped[row["榜单类型"]].append(row)
    for items in dimension_grouped.values():
        items.sort(key=lambda item: as_int(item["榜内排名"]))

    profile_type_counts = Counter(row["机构类型"] for row in profile_rows)
    top100_type_counts = Counter(row["机构类型"] for row in type_review)
    top100_level_counts = Counter(row["机构层级"] for row in type_review)

    foreign_top20 = [
        row["标准化机构名称"]
        for row in top20
        if row["主属国家/地区"] != "中国"
    ]

    return {
        "merged_rows": merged_rows,
        "top20": top20,
        "weights": weights,
        "dimension_grouped": dimension_grouped,
        "type_review": type_review,
        "preprocess_rows": preprocess_rows,
        "selected_models": selected_models,
        "strength_top3": strength_top3,
        "evolution_rows": evolution_rows,
        "type_grouped": type_grouped,
        "level_grouped": level_grouped,
        "profile_rows": profile_rows,
        "profile_type_counts": profile_type_counts,
        "top100_type_counts": top100_type_counts,
        "top100_level_counts": top100_level_counts,
        "foreign_top20": foreign_top20,
    }


def build_metrics(data: dict[str, object]) -> dict[str, object]:
    preprocess_rows = data["preprocess_rows"]
    selected_models = data["selected_models"]
    top20 = data["top20"]
    profile_type_counts: Counter[str] = data["profile_type_counts"]
    top100_type_counts: Counter[str] = data["top100_type_counts"]
    top100_level_counts: Counter[str] = data["top100_level_counts"]
    strength_top3 = data["strength_top3"]
    type_grouped = data["type_grouped"]
    level_grouped = data["level_grouped"]
    dimension_grouped = data["dimension_grouped"]

    total_docs = len(data["merged_rows"])
    total_model_docs = sum(as_int(row["进入建模文献数"]) for row in preprocess_rows)
    periods = [row["阶段"] for row in preprocess_rows]
    top10 = top20[:10]
    top10_names = [row["标准化机构名称"] for row in top10]

    selected_k = {
        stage: as_int(row["候选主题数"])
        for stage, row in selected_models.items()
    }

    example_paths = []
    preferred_path_starts = (
        "储层 / 孔隙 / 渗透率 / 孔隙度",
        "催化剂 / 地层 / 产品 / 收率",
        "储层 / 驱替 / 注入 / 采收率",
        "二氧化碳 / 实验 / 模拟 / 储层",
    )
    seen_paths: set[str] = set()
    for start_label in preferred_path_starts:
        for row in data["evolution_rows"]:
            labels = [
                row["2011-2015主题标签"],
                row["2016-2020主题标签"],
                row["2021-2025主题标签"],
            ]
            if labels[0] != start_label:
                continue
            path_text = " → ".join(label for label in labels if label)
            if path_text and path_text not in seen_paths:
                example_paths.append(path_text)
                seen_paths.add(path_text)
                break

    return {
        "cover_info": load_cover_info(),
        "total_docs": total_docs,
        "total_model_docs": total_model_docs,
        "periods": periods,
        "top10": top10,
        "top10_names": top10_names,
        "foreign_top20": data["foreign_top20"],
        "profile_type_counts": profile_type_counts,
        "top100_type_counts": top100_type_counts,
        "top100_level_counts": top100_level_counts,
        "selected_models": selected_models,
        "selected_k": selected_k,
        "strength_top3": strength_top3,
        "type_grouped": type_grouped,
        "level_grouped": level_grouped,
        "dimension_grouped": dimension_grouped,
        "example_paths": example_paths[:3],
    }


def build_abstract_zh(metrics: dict[str, object], data: dict[str, object]) -> str:
    top1 = metrics["top10"][0]
    top2 = metrics["top10"][1]
    profile_type_counts: Counter[str] = metrics["profile_type_counts"]
    return (
        "本文以已筛定的109本高质量期刊石油领域文献数据为基础数据源，整合CNKI、WOS与CSCD三库"
        f"2011-2025年记录，构建覆盖{metrics['total_docs']}篇文献的双主键去重正式语料。研究首先按照"
        "“标准DOI键优先、元数据回退键补充”的规则完成跨库去重，并通过机构译名清洗、标准化归并与子单元过滤"
        "构建机构分析底表；随后从科研产出、学术影响、合作与国际化三个维度提取7项指标，采用先验权重、熵权法"
        "与CRITIC法的组合赋权方案，通过TOPSIS对全球石油领域核心研究机构开展综合评价；最后基于"
        f"{metrics['total_model_docs']}篇建模文献，按2011-2015、2016-2020、2021-2025三个阶段构建LDA主题模型，"
        "分析研究热点演化路径以及不同机构类型、机构层级的主题偏好差异。结果表明：TOPSIS综合排名前两位分别为"
        f"{top1['标准化机构名称']}和{top2['标准化机构名称']}，其中{top1['标准化机构名称']}综合得分为"
        f"{top1['TOPSIS综合得分']}；机构画像环节共识别企业研发中心{profile_type_counts.get('企业研发中心', 0)}家、"
        f"高校/科研院所{profile_type_counts.get('高校/科研院所', 0)}家、政府机构{profile_type_counts.get('政府机构', 0)}家；"
        "三阶段模型均选取11个主题，研究热点总体呈现出由盆地构造、油田开发与传统催化转化，逐步转向储层预测、"
        "低碳利用与高精度催化表征的趋势。高校/科研院所在知识探索型主题上更具优势，企业研发中心则更集中于"
        "工程开发与应用导向主题。研究为石油领域科研机构竞争力诊断、主题布局优化和资源配置决策提供了可复现的"
        "数据分析路径。"
    )


def build_abstract_en(metrics: dict[str, object]) -> str:
    top1 = ENGLISH_NAME_MAP.get(metrics["top10"][0]["标准化机构名称"], metrics["top10"][0]["标准化机构名称"])
    top2 = ENGLISH_NAME_MAP.get(metrics["top10"][1]["标准化机构名称"], metrics["top10"][1]["标准化机构名称"])
    profile_type_counts: Counter[str] = metrics["profile_type_counts"]
    return (
        "Using literature records derived from 109 high-quality petroleum journals, this study "
        "integrates CNKI, Web of Science, and CSCD data from 2011 to 2025 and constructs a dual-key "
        f"deduplicated formal corpus containing {metrics['total_docs']} papers. The workflow first "
        "implements cross-database deduplication with standardized DOI keys and metadata fallback keys, "
        "and then performs institution translation cleaning, normalization, and sub-unit filtering to "
        "build an institution-level analytic dataset. Based on three dimensions of research output, "
        "academic impact, and collaboration/internationalization, seven indicators are retained and a "
        "combined weighting scheme averaging prior weights, entropy weights, and CRITIC weights is used "
        "for TOPSIS evaluation. A three-stage LDA framework is further employed to examine thematic "
        f"evolution using {metrics['total_model_docs']} model-ready documents. The results show that "
        f"{top1} and {top2} rank first and second in the comprehensive TOPSIS evaluation. Institution "
        f"profiling identifies {profile_type_counts.get('企业研发中心', 0)} enterprise R&D centers, "
        f"{profile_type_counts.get('高校/科研院所', 0)} universities/research institutes, and "
        f"{profile_type_counts.get('政府机构', 0)} government institutions. Across the three periods, "
        "the optimal topic number remains 11, and the thematic focus shifts from basin structure, oilfield "
        "development, and conventional catalytic conversion toward reservoir prediction, low-carbon "
        "utilization, and refined catalytic active-site studies. Universities and research institutes are "
        "more prominent in exploratory knowledge-intensive topics, whereas enterprise R&D centers are more "
        "concentrated in engineering and application-oriented themes. This study provides a reproducible "
        "analytical route for evaluating institutional competitiveness and understanding topic evolution in "
        "the global petroleum research field."
    )


def build_tables(metrics: dict[str, object], data: dict[str, object]) -> dict[str, str]:
    weights_table = markdown_table(
        ["指标名称", "所属维度", "组合权重"],
        [
            [
                row["label_zh"],
                row["dimension_zh"],
                f"{as_float(row['combined_weight']):.6f}",
            ]
            for row in data["weights"]
        ],
    )

    top10_table = markdown_table(
        ["排名", "机构名称", "国家/地区", "综合得分", "科研产出子得分", "学术影响子得分", "合作与国际化子得分"],
        [
            [
                row["TOPSIS综合排名"],
                row["标准化机构名称"],
                row["主属国家/地区"],
                row["TOPSIS综合得分"],
                row["科研产出子得分"],
                row["学术影响子得分"],
                row["合作与国际化子得分"],
            ]
            for row in metrics["top10"]
        ],
    )

    model_table = markdown_table(
        ["阶段", "阶段文献数", "通过领域过滤文献数", "进入建模文献数", "选定主题数", "困惑度", "一致性得分"],
        [
            [
                row["阶段"],
                row["阶段文献数"],
                row["通过领域过滤文献数"],
                row["进入建模文献数"],
                str(metrics["selected_k"][row["阶段"]]),
                metrics["selected_models"][row["阶段"]]["困惑度"],
                metrics["selected_models"][row["阶段"]]["一致性得分"],
            ]
            for row in data["preprocess_rows"]
        ],
    )

    profile_table = markdown_table(
        ["机构类型", "机构数量", "说明"],
        [
            ["企业研发中心", str(metrics["profile_type_counts"].get("企业研发中心", 0)), "标准化后数量最多，反映产业主体广泛参与"],
            ["高校/科研院所", str(metrics["profile_type_counts"].get("高校/科研院所", 0)), "在高影响与头部排名中占据主导"],
            ["政府机构", str(metrics["profile_type_counts"].get("政府机构", 0)), "样本较少，仅作为补充观察"],
            ["其他", str(metrics["profile_type_counts"].get("其他", 0)), "未纳入正式类型比较"],
        ],
    )

    return {
        "weights_table": weights_table,
        "top10_table": top10_table,
        "model_table": model_table,
        "profile_table": profile_table,
    }


def build_markdown(metrics: dict[str, object], data: dict[str, object]) -> str:
    tables = build_tables(metrics, data)
    strength_top3 = metrics["strength_top3"]
    type_grouped = metrics["type_grouped"]
    level_grouped = metrics["level_grouped"]
    dimension_grouped = metrics["dimension_grouped"]
    foreign_text = join_cn(metrics["foreign_top20"])
    top10_names_text = join_cn(metrics["top10_names"])

    lines: list[str] = []
    lines.append(f"# {TITLE_ZH}")
    lines.append("")
    lines.append("## 摘要")
    lines.append(build_abstract_zh(metrics, data))
    lines.append("")
    lines.append("**关键词：**" + "；".join(KEYWORDS_ZH))
    lines.append("")
    lines.append("## Abstract")
    lines.append(f"**{TITLE_EN}**")
    lines.append("")
    lines.append(build_abstract_en(metrics))
    lines.append("")
    lines.append("**Key Words:** " + "; ".join(KEYWORDS_EN))
    lines.append("")
    lines.append("## 目录")
    lines.append("摘要")
    lines.append("Abstract")
    lines.append("1 绪论")
    lines.append("2 数据来源与研究方法")
    lines.append("3 全球石油领域研究机构学术影响力评价结果")
    lines.append("4 研究主题演化与机构画像分析")
    lines.append("5 讨论与建议")
    lines.append("6 结论")
    lines.append("参考文献")
    lines.append("致谢")
    lines.append("")

    lines.append("## 1 绪论")
    lines.append("")
    lines.append("### 1.1 研究背景与意义")
    lines.append(
        "石油仍然是全球能源结构与工业原料体系中的关键支柱。随着非常规油气开发、储层精细表征、"
        "二氧化碳利用与封存、催化转化等议题持续升温，科研机构在石油领域中的学术地位、知识生产方式"
        "和研究主题布局，已经成为观察行业科技竞争格局的重要窗口。与通用性的大学排名不同，石油领域"
        "机构评价不仅涉及论文数量和被引表现，还受到合作网络、工程应用导向以及垂直领域主题契合度的"
        "共同影响。"
    )
    lines.append(
        "从现实需求看，面向全球石油科技竞争与能源安全治理，单纯依靠经验判断已难以支撑机构竞争力"
        "识别和资源配置决策。以高质量期刊文献为基础，综合评价全球石油领域研究机构的学术影响力，并"
        "进一步追踪热点主题的演化轨迹，有助于识别头部机构优势来源、发现特色机构突破方向，也有助于"
        "为高校、企业研发主体和管理部门提供差异化决策参考。"
    )
    lines.append("")
    lines.append("### 1.2 国内外研究现状")
    lines.append(
        "现有研究大致可分为两条主线：一条聚焦科研机构或学者学术影响力评价，强调多指标综合评价模型"
        "的构建；另一条聚焦主题识别与主题演化分析，强调从大规模文本中抽取研究热点及其变化路径。"
        "Schlögl等指出，机构评价结果对数据源选择、计数方式和指标属性设置高度敏感，评价框架的透明性"
        "与可解释性是保证结果可信的前提[1]。在国内，林子婕、唐星龙等从多维互动视角讨论学术影响力评价，"
        "张璜、高睿则从主题性视角引入LDA模型分析学术影响力问题[3-4]。"
    )
    lines.append(
        "在石油及相关领域研究中，已有成果更多集中于期刊影响力、单一高校或单一机构群体的科学计量分析，"
        "例如针对石油天然气工程类期刊影响力的研究[6]，以及围绕能源类高校研究生产力开展的实证分析[7-8]。"
        "这些研究为垂直领域评价提供了经验，但仍存在三个不足：一是评价对象往往局限于少量机构或单一国家，"
        "缺乏全球比较；二是机构评价与主题演化通常分离处理，尚未形成从影响力到主题布局的贯通解释；三是"
        "对于多源异构数据清洗、跨库去重和机构标准化的规则披露不充分，导致研究复现难度较高。"
    )
    lines.append(
        "主题演化研究方面，近年来文献计量与文本挖掘方法结合愈发紧密。王宏宇等基于语义网络识别科研主题"
        "演化路径[5]，Blei等提出的LDA模型则为主题发现提供了经典技术框架[12]。然而，面向全球石油领域"
        "研究机构的主题演化分析，若缺乏稳定的机构标准化和类型划分，就难以进一步回答“何种机构在何种主题"
        "上形成优势、其优势又如何反馈到综合影响力”这一关键问题。"
    )
    lines.append("")
    lines.append("### 1.3 研究问题与研究目标")
    lines.append(
        "围绕上述不足，本文重点回答三个问题。第一，如何在高质量期刊文献约束下，构建覆盖2011-2025年的"
        "全球石油领域正式分析语料，并形成可在论文中合理表述的数据清洗与跨库去重规则。第二，如何在"
        "现有字段条件下，建立兼顾科研产出、学术影响和合作国际化的机构综合评价模型，并据此识别全球"
        "石油领域核心研究机构的层级结构。第三，如何将机构评价结果与主题演化结果联系起来，解释不同"
        "机构类型和机构层级的主题偏好及其变化。"
    )
    lines.append(
        "据此，本文的研究目标包括：构建双主键正式语料并完成机构标准化；形成可复现的TOPSIS机构影响力"
        "评价结果与层级划分；在三阶段LDA框架下识别石油领域主题演化路径；比较不同机构类型与机构层级"
        "的主题差异，并在此基础上提出具有现实针对性的管理建议。"
    )
    lines.append("")
    lines.append("### 1.4 研究思路与可能创新")
    lines.append(
        "本文将“数据治理—机构评价—主题演化—策略建议”串联为一条完整分析链条。首先，在原始多源文献"
        "基础上执行源内可用性清洗和跨库双主键去重；其次，通过机构译名纠偏、母机构归并和子单元过滤"
        "构建标准化机构主表；再次，在Top100核心机构样本上实施组合赋权与TOPSIS评价，并建立机构类型和"
        "机构层级画像；最后，复用正式语料对应的主题分析结果，从时间、类型和层级三个维度解释石油领域"
        "研究主题的演化逻辑。"
    )
    lines.append(
        "本文的主要创新点体现在三个方面：其一，在论文初稿中显式给出数据清洗、去重与机构标准化规则，"
        "增强结果可复查性；其二，将机构影响力评价与主题演化分析放在同一正式语料和统一机构口径下讨论，"
        "避免传统研究中“评价结果”与“主题分析”各自为政的割裂问题；其三，在开题报告原计划框架基础上，"
        "根据实际数据可得性对指标体系和类型比较边界进行了收敛，使论文结论更具科学性和可辩护性。"
    )
    lines.append("")

    lines.append("## 2 数据来源与研究方法")
    lines.append("")
    lines.append("### 2.1 数据来源与研究范围")
    lines.append(
        f"本文将当前持有的三库原始数据统一视为来源于已筛定的109本高质量期刊石油领域文献数据，并以此作为"
        f"研究的基础数据源。正式分析语料来自CNKI、WOS与CSCD三库2011-2025年记录，经统一清洗和双主键跨库"
        f"去重后形成正式主表，共包含{metrics['total_docs']}篇文献。按主题建模统计口径，2011-2015、"
        f"2016-2020、2021-2025三个阶段的文献数分别为{data['preprocess_rows'][0]['阶段文献数']}篇、"
        f"{data['preprocess_rows'][1]['阶段文献数']}篇和{data['preprocess_rows'][2]['阶段文献数']}篇。"
        "由于本文研究对象是研究机构而非期刊本身，高质量期刊属性在研究中主要承担数据来源约束作用，而不作为"
        "后续评价指标。"
    )
    lines.append("")
    lines.append("### 2.2 数据清洗与跨库去重规则")
    lines.append(
        "为保证语料构建过程可重复、可审查且可在论文中清晰表述，本文将数据清洗与跨库去重划分为“源内可用性"
        "清洗”与“跨库文献标识合并”两个阶段。在源内清洗阶段，CNKI记录不再以DOI是否存在作为入库门槛，而是"
        "要求题名、发表时间、期刊和机构四类核心书目信息至少具备可用字段，以避免因中文数据库DOI缺失而无必要"
        "地剔除仍具识别价值的文献。"
    )
    lines.append(
        "在跨库去重阶段，本文采用“双主键”合并策略。第一主键为严格标准DOI键，统一从doi和registered_doi"
        "字段中提取；第二主键为由规范化题名、年份、第一作者和期刊共同构成的元数据回退键。具体规则为：当两条"
        "记录共享同一可用标准DOI时判定为同一文献；当记录缺乏可用标准DOI时，仅在元数据键唯一对应既有分组且"
        "不存在DOI冲突的前提下才允许执行元数据合并；若同一元数据键对应多个候选分组，则视为歧义匹配并保留到"
        "审查表，而不自动吞并。"
    )
    lines.append(
        "完成分组合并后，本文采用来源感知的字段择优融合策略生成正式主表。中文题名、中文摘要和中文关键词优先"
        "保留CNKI或CSCD信息，英文题名、英文摘要、英文关键词和被引频次优先保留WOS信息，同时显式保留来源库组合、"
        "去重依据、标准DOI键和元数据去重键等字段，以支持后续结果追踪。该规则兼顾了语料覆盖度和跨库误并风险控制。"
    )
    lines.append("")
    lines.append("### 2.3 机构标准化与研究样本构建")
    lines.append(
        f"机构标准化遵循“译名可读、标准名统一、颗粒度一致”的原则。对实验室、学院、研究中心、分院等子单元，"
        "若能够稳定回溯到母机构，则统一归并到母机构层级；对具有独立学术身份且无法合理归并的机构，则保留独立"
        "标准名。经清洗后，机构画像表共保留7882个有效标准化机构，其中企业研发中心"
        f"{metrics['profile_type_counts'].get('企业研发中心', 0)}个，高校/科研院所"
        f"{metrics['profile_type_counts'].get('高校/科研院所', 0)}个，政府机构"
        f"{metrics['profile_type_counts'].get('政府机构', 0)}个，另有"
        f"{metrics['profile_type_counts'].get('其他', 0)}个机构暂列“其他”，不纳入正式类型比较。"
    )
    lines.append(
        "在核心评价样本方面，本文依据标准化机构频次表构建Top100核心机构集合。类型复核结果显示，Top100样本中"
        f"高校/科研院所{metrics['top100_type_counts'].get('高校/科研院所', 0)}家、企业研发中心"
        f"{metrics['top100_type_counts'].get('企业研发中心', 0)}家、政府机构"
        f"{metrics['top100_type_counts'].get('政府机构', 0)}家；按层级划分则包括头部引领型"
        f"{metrics['top100_level_counts'].get('头部引领型', 0)}家、中坚创新型"
        f"{metrics['top100_level_counts'].get('中坚创新型', 0)}家和特色细分型"
        f"{metrics['top100_level_counts'].get('特色细分型', 0)}家。"
    )
    lines.append("")
    lines.append("表2-1 机构类型分布情况")
    lines.append("")
    lines.append(tables["profile_table"])
    lines.append("")
    lines.append("### 2.4 指标体系与组合赋权方法")
    lines.append(
        "结合开题报告中“科研产出—学术影响—国际合作”三维框架与当前数据可得性，本文最终保留7项指标进入综合评价"
        "模型，即去重论文总数、近五年发文占比、H指数、高被引论文占比、合作论文占比、国际合作论文占比和合作国家/"
        "地区数。由于当前重建过程未恢复出可追溯的专家评分数据，本文不再将德尔菲法作为实际赋权手段，而是采用"
        "“维度均衡先验权重+熵权法+CRITIC法”的组合赋权方案。三类权重先分别计算，再进行算术平均和归一化，形成"
        "最终组合权重，以兼顾先验平衡、样本离散度和指标冲突性。"
    )
    lines.append("")
    lines.append("表2-2 机构评价指标组合权重")
    lines.append("")
    lines.append(tables["weights_table"])
    lines.append("")
    lines.append(
        "从组合权重看，去重论文总数权重最高，为0.279947，说明样本规模仍是机构综合影响力的基础；国际合作论文占比"
        "与近五年发文占比分别承担国际化活跃度和近期活跃度识别功能；H指数、高被引论文占比则主要反映成果影响质量。"
    )
    lines.append("")
    lines.append("### 2.5 TOPSIS评价与机构分层")
    lines.append(
        "在综合评价阶段，本文先对全部正向指标进行标准化处理，再依据组合权重构建加权标准化决策矩阵；随后确定正理想"
        "解与负理想解，计算各机构到两类理想解的距离，并据此获得TOPSIS综合贴近度得分。得分越高，表示机构越接近"
        "理想最优状态。基于综合得分分布，进一步将Top100机构划分为头部引领型、中坚创新型和特色细分型三个层级，"
        "用于后续机构画像和主题偏好比较。"
    )
    lines.append("")
    lines.append("### 2.6 主题演化分析流程")
    lines.append(
        f"主题分析以正式语料中的摘要文本为对象，按2011-2015、2016-2020、2021-2025三个阶段分别建模。"
        "文本预处理主要包括摘要筛选、分词、去停用词、石油领域术语归并与词汇表截断。主题数选择采用困惑度与一致性"
        "得分联合判定，三个阶段最终均选择11个主题。随后，本文依据相邻阶段主题词分布相似度构建演化连接，仅保留"
        "满足阈值要求的路径，以识别主题的延续、分化与迁移关系。"
    )
    lines.append("")
    lines.append("表2-3 三阶段主题建模与选模结果")
    lines.append("")
    lines.append(tables["model_table"])
    lines.append("")

    lines.append("## 3 全球石油领域研究机构学术影响力评价结果")
    lines.append("")
    lines.append("### 3.1 样本概况与综合格局")
    lines.append(
        f"依据Top100核心机构综合评价结果，2011-2025年全球石油领域研究机构影响力排名前十位依次为"
        f"{top10_names_text}。综合排名呈现出明显的“头部集中+多元并存”格局：中国科学院和中国石油大学（北京）"
        "构成第一梯队，其后依次为两大石油央企主体以及行业特色高校与综合性大学。"
    )
    lines.append(
        f"综合排名前20位中除国内头部高校与科研机构外，还出现了{foreign_text}等境外高水平大学。说明在石油领域的"
        "高质量文献场域中，国际合作广度和学术影响质量能够帮助部分境外机构在总体发文规模不占绝对优势的情况下"
        "进入前列。"
    )
    lines.append("")
    lines.append("表3-1 TOPSIS综合排名前十机构")
    lines.append("")
    lines.append(tables["top10_table"])
    lines.append("")
    lines.append("### 3.2 TOPSIS综合结果分析")
    lines.append(
        f"从综合得分看，中国科学院以{metrics['top10'][0]['TOPSIS综合得分']}位列第一，其优势来源于高水平的科研产出"
        "和稳健的学术影响表现；中国石油大学（北京）以"
        f"{metrics['top10'][1]['TOPSIS综合得分']}位列第二，在发文规模与持续活跃度上保持显著优势。"
        "排名第三和第四的中国石油化工股份有限公司、中国石油天然气股份有限公司表明，大型行业企业不仅在工程应用"
        "层面发挥关键作用，也已经形成持续稳定的高质量科研输出能力。"
    )
    lines.append(
        "进一步看，头部机构的优势并不完全相同。中国科学院表现为多维均衡型优势；中国石油大学（北京）体现出"
        "主题契合度与长期积累形成的产出优势；中国石油化工股份有限公司和中国石油天然气股份有限公司则显示出"
        "企业研发主体在行业问题导向研究上的持续发力。"
    )
    lines.append("")
    lines.append("### 3.3 分维度优势比较")
    lines.append(
        f"从科研产出维度看，前五位分别为{join_cn([row['标准化机构名称'] for row in dimension_grouped['科研产出维度'][:5]])}；"
        f"从学术影响维度看，前五位分别为{join_cn([row['标准化机构名称'] for row in dimension_grouped['学术影响维度'][:5]])}；"
        f"从合作与国际化维度看，前五位分别为{join_cn([row['标准化机构名称'] for row in dimension_grouped['合作与国际化维度'][:5]])}。"
    )
    lines.append(
        "这一结果说明，机构综合影响力并不完全等同于单一维度优势。科研产出维度主要受发文规模驱动，国内头部科研院所"
        "和行业高校占据主导；学术影响维度中，湖南大学、新加坡国立大学、华南理工大学等机构表现突出，体现出“质量"
        "优先型”特征；合作与国际化维度中，帝国理工学院、南洋理工大学、昆士兰大学等机构优势显著，反映出国际合作"
        "网络对综合评价结果的支撑作用。"
    )
    lines.append("")
    lines.append("### 3.4 机构类型与机构层级特征")
    lines.append(
        "从全部标准化机构画像看，企业研发中心在数量上占绝对多数，但在Top100高影响样本中，高校/科研院所占比更高，"
        "说明在高质量期刊口径下，知识生产质量与学术扩散能力仍主要由高校与科研机构主导。与此同时，少数石油央企和"
        "大型企业研发主体能够凭借行业问题驱动的稳定产出进入头部样本，体现出鲜明的产业技术牵引特征。"
    )
    lines.append(
        "按机构层级观察，头部引领型机构在综合得分、合作网络和成果质量等方面表现更为均衡；中坚创新型机构往往在"
        "某一维度形成差异化突破；特色细分型机构则更多依赖细分主题长期积累进入核心样本。该分层结果为后续主题偏好"
        "分析提供了稳定的机构画像基础。"
    )
    lines.append("")

    lines.append("## 4 研究主题演化与机构画像分析")
    lines.append("")
    lines.append("### 4.1 三阶段建模结果与热点主题")
    lines.append(
        f"三阶段LDA建模结果显示，各阶段最优主题数均为11。2011-2015年主题强度最高的三个主题分别为"
        f"{join_cn([row['主题标签'] for row in strength_top3['2011-2015']])}；2016-2020年对应为"
        f"{join_cn([row['主题标签'] for row in strength_top3['2016-2020']])}；2021-2025年对应为"
        f"{join_cn([row['主题标签'] for row in strength_top3['2021-2025']])}。"
    )
    lines.append(
        "从整体趋势看，石油领域研究热点经历了由传统盆地构造、油田开发和催化转化问题向储层预测、低碳利用与高精度"
        "催化位点表征逐步演化的过程。也就是说，研究重心并非简单替换，而是在传统资源开发议题基础上不断叠加储层"
        "精细表征、绿色转型和材料微观机制等新方向。"
    )
    lines.append("")
    lines.append(
        f"![图4-1 主题强度演化曲线]({md_image_path(TOPIC_DIR / 'topic_intensity_evolution_curve_dual_key_2011_2025.png')})"
    )
    lines.append("")
    lines.append("### 4.2 主题演化路径分析")
    lines.append(
        f"在相邻阶段主题相似度匹配基础上，共识别出{len(data['evolution_rows'])}条主题演化路径。较具代表性的路径包括："
        + "；".join(metrics["example_paths"])
        + "。"
    )
    lines.append(
        "上述路径表明，石油领域主题演化具有明显的延续性与细化性。一方面，储层表征和开发类研究沿着“宏观构造—储层"
        "性质—微观孔隙结构”的路线不断深化；另一方面，催化研究则沿着“反应现象—材料活性—位点机理”的方向不断"
        "细分。二氧化碳相关主题的持续出现，进一步说明绿色低碳议题已经成为石油科技研究的重要增长点。"
    )
    lines.append("")
    lines.append("### 4.3 机构类型主题偏好比较")
    lines.append(
        f"类型比较结果显示，2011-2015年高校/科研院所占比最高的主题为"
        f"{join_cn([row['主题标签'] for row in type_grouped[('2011-2015', '高校/科研院所')][:2]])}，"
        f"企业研发中心则更集中于{join_cn([row['主题标签'] for row in type_grouped[('2011-2015', '企业研发中心')][:2]])}。"
        f"2016-2020年，高校/科研院所主要集中于"
        f"{join_cn([row['主题标签'] for row in type_grouped[('2016-2020', '高校/科研院所')][:2]])}，"
        f"企业研发中心主要集中于{join_cn([row['主题标签'] for row in type_grouped[('2016-2020', '企业研发中心')][:2]])}。"
        f"到2021-2025年，高校/科研院所更偏向"
        f"{join_cn([row['主题标签'] for row in type_grouped[('2021-2025', '高校/科研院所')][:2]])}，"
        f"企业研发中心则集中于{join_cn([row['主题标签'] for row in type_grouped[('2021-2025', '企业研发中心')][:2]])}。"
    )
    lines.append(
        "由此可见，高校/科研院所更容易在知识探索和前沿主题上保持分布广度，而企业研发中心更偏向与油田开发、储层"
        "预测和工程应用直接相关的问题。政府机构样本仅2家，虽然在部分阶段也呈现出与低碳利用、催化转化相关的主题"
        "聚集，但由于样本规模过小，本文仅作补充观察而不展开正式比较。"
    )
    lines.append("")
    lines.append(
        f"![图4-2 不同机构类型主题分布热力图]({md_image_path(TOPIC_DIR / 'topic_distribution_heatmap_by_institution_type_dual_key_2011_2025.png')})"
    )
    lines.append("")
    lines.append("### 4.4 机构层级主题偏好比较")
    lines.append(
        f"从机构层级看，2011-2015年头部引领型机构的代表主题为"
        f"{level_grouped[('2011-2015', '头部引领型')][0]['主题标签']}，中坚创新型机构更偏向"
        f"{level_grouped[('2011-2015', '中坚创新型')][0]['主题标签']}；2016-2020年，头部引领型与特色细分型机构"
        f"都在{level_grouped[('2016-2020', '头部引领型')][0]['主题标签']}等储层表征主题上保持活跃，而中坚创新型"
        f"机构更多集中于{level_grouped[('2016-2020', '中坚创新型')][0]['主题标签']}；2021-2025年，头部引领型机构"
        f"强化了{level_grouped[('2021-2025', '头部引领型')][0]['主题标签']}等高复杂度主题，中坚创新型机构则更偏向"
        f"{level_grouped[('2021-2025', '中坚创新型')][0]['主题标签']}。"
    )
    lines.append(
        "总体而言，机构层级越高，越容易同时覆盖基础研究、工程预测和国际合作关联主题；层级越细分的机构，则越倾向于"
        "在少数技术主题上形成高占比优势。这说明影响力层级不仅是结果性排名，也可视为主题布局广度和资源整合能力的"
        "一个侧面反映。"
    )
    lines.append("")
    lines.append(
        f"![图4-3 不同机构层级主题分布热力图]({md_image_path(TOPIC_DIR / 'topic_distribution_heatmap_by_institution_level_dual_key_2011_2025.png')})"
    )
    lines.append("")

    lines.append("## 5 讨论与建议")
    lines.append("")
    lines.append("### 5.1 机构影响力与主题布局的耦合关系")
    lines.append(
        "综合前文结果可以发现，石油领域机构影响力并非简单由规模决定，而是由规模、质量、合作和主题布局共同塑造。"
        "头部机构往往既拥有较强的科研产出能力，也能覆盖更多前沿主题和国际合作网络；中坚机构更容易在单一维度上"
        "形成突破；特色机构则更多通过细分技术路线保持存在感。换言之，机构影响力评价结果与主题演化结果之间存在"
        "较强的结构性对应关系。"
    )
    lines.append("")
    lines.append("### 5.2 对高校/科研院所的建议")
    lines.append(
        "高校/科研院所应继续保持在储层精细表征、低碳利用、催化机理等知识探索型主题上的优势，同时加强与企业场景的"
        "深度耦合，推动基础研究问题向工程可验证问题延伸。对于已进入头部引领型的机构，应重点提升跨机构协作效率和"
        "国际合作深度；对于中坚创新型机构，则应在已有优势主题上形成更清晰的学术品牌。"
    )
    lines.append("")
    lines.append("### 5.3 对企业研发中心的建议")
    lines.append(
        "企业研发中心在油田开发、储层预测与工程应用类主题上已形成明显优势，但在高影响成果产出和国际合作方面仍有"
        "提升空间。建议企业主体在保持问题导向优势的同时，进一步布局催化材料、二氧化碳利用与封存等前沿议题，强化"
        "与高校、科研院所的协同研发机制，从而提升研究成果的学术可见度与长期影响力。"
    )
    lines.append("")
    lines.append("### 5.4 对管理部门的建议")
    lines.append(
        "管理部门应推动高质量文献数据、机构标准名和主题监测结果的常态化归集，建立面向重点能源技术方向的滚动监测"
        "机制。对于跨机构协同创新，可考虑围绕储层预测、低碳利用和关键催化材料等主题建立联合攻关网络，促进基础"
        "研究资源和产业场景资源的有效对接。"
    )
    lines.append("")
    lines.append("### 5.5 研究局限")
    lines.append(
        f"本文仍存在三点局限。第一，机构类型识别虽已显著改善，但仍有{metrics['profile_type_counts'].get('其他', 0)}个"
        "机构暂列“其他”，说明少量跨国联盟、中心类组织的属性辨识仍需进一步细化。第二，由于缺乏可追溯的专家评分"
        "数据，本文未将德尔菲法纳入实际赋权流程，而是采用完全可复现的组合客观赋权方案。第三，经有效机构清洗后，"
        "国际组织未形成稳定样本，因此机构类型比较仅正式保留高校/科研院所、企业研发中心和政府机构三类，并将国际"
        "组织样本不足作为明确限制说明。"
    )
    lines.append("")

    lines.append("## 6 结论")
    lines.append("")
    lines.append(
        f"本文基于109本高质量期刊石油领域文献数据，构建了覆盖2011-2025年的双主键正式语料，并完成了从机构标准化、"
        "Top100综合评价到主题演化比较的一体化分析。研究表明，中国科学院、中国石油大学（北京）等机构在综合评价中"
        "位居前列，头部格局体现出大型科研院所、行业特色高校与石油企业主体并存的结构特征。"
    )
    lines.append(
        "主题演化结果显示，石油领域研究热点总体呈现由盆地构造与传统开发问题向储层预测、低碳利用和高精度催化机理"
        "演化的趋势。高校/科研院所在知识探索型主题上保持优势，企业研发中心则更集中于工程应用与开发类议题，机构"
        "层级越高的样本越容易表现出更宽的主题覆盖度。"
    )
    lines.append(
        "总体而言，本文为高质量文献驱动下的全球石油领域研究机构评价提供了可复现的技术路线，也为后续机构竞争力诊断、"
        "主题监测和资源配置优化提供了数据基础。"
    )
    lines.append("")

    lines.append("## 参考文献")
    lines.append("")
    for index, ref in enumerate(REFERENCES, start=1):
        lines.append(f"[{index}] {ref}")
    lines.append("")
    lines.append("## 致谢")
    lines.append("")
    lines.append(
        "感谢导师在选题、方法与写作过程中的指导，感谢学院提供的研究条件与毕业论文组织支持。本文当前版本为基于"
        "最新重建数据链生成的论文初稿，后续仍将结合参考文献细化、正文扩写、图表编号校核和版式套版结果继续完善。"
    )
    lines.append("")
    return "\n".join(lines)


def set_run_font(run, size: float = 12, bold: bool = False, east_asia: str = "宋体", ascii_font: str = "Times New Roman") -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_common(paragraph, first_line_indent: float = 0.85, align: int = WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25


def add_text_paragraph(document: Document, text: str, *, center: bool = False, bold: bool = False, size: float = 12, east_asia: str = "宋体", first_line_indent: float = 0.85) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_common(
        paragraph,
        first_line_indent=0 if center else first_line_indent,
        align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, east_asia=east_asia)


def add_heading_1(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True, east_asia="黑体")


def add_heading_2(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_common(paragraph, first_line_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True, east_asia="黑体")


def add_heading_3(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_common(paragraph, first_line_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True, east_asia="黑体")


def add_table_title(document: Document, text: str) -> None:
    add_text_paragraph(document, text, center=True, size=11)


def add_table_docx(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        paragraph = table.rows[0].cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_run_font(run, size=10.5)
    add_text_paragraph(document, "资料来源：作者根据重建结果表整理。", size=10.5, first_line_indent=0)


def add_figure(document: Document, path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(15.5))
    add_text_paragraph(document, caption, center=True, size=10.5, first_line_indent=0)
    add_text_paragraph(document, "资料来源：作者根据重建结果图整理。", center=True, size=10.5, first_line_indent=0)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=10.5)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.different_first_page_header_footer = True

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_paragraph.add_run("中国石油大学（北京）本科毕业论文")
    set_run_font(header_run, size=10.5)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer_paragraph)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)


def add_cover_page(document: Document, cover_info: dict[str, str]) -> None:
    add_text_paragraph(document, "中国石油大学（北京）本科毕业论文", center=True, bold=True, size=18, east_asia="黑体")
    add_text_paragraph(document, "论文初稿", center=True, size=14, east_asia="黑体")
    add_text_paragraph(document, "", center=True)
    add_text_paragraph(document, TITLE_ZH, center=True, bold=True, size=16, east_asia="黑体")
    add_text_paragraph(document, TITLE_EN, center=True, size=12, east_asia="Times New Roman", first_line_indent=0)
    add_text_paragraph(document, "", center=True)
    add_text_paragraph(document, f"学生姓名：{cover_info.get('姓名', '')}", center=True, size=12)
    add_text_paragraph(document, f"学号：{cover_info.get('学号', '')}", center=True, size=12)
    add_text_paragraph(document, f"学院：{cover_info.get('学院', '')}", center=True, size=12)
    add_text_paragraph(document, f"专业：{cover_info.get('学科专业', '')}", center=True, size=12)
    add_text_paragraph(document, "完成日期：2026年4月24日", center=True, size=12)
    document.add_page_break()


def add_docx_content(document: Document, markdown_text: str, metrics: dict[str, object], data: dict[str, object]) -> None:
    add_heading_1(document, "摘要")
    add_text_paragraph(document, build_abstract_zh(metrics, data))
    add_text_paragraph(document, "关键词：" + "；".join(KEYWORDS_ZH), first_line_indent=0)

    add_heading_1(document, "Abstract")
    add_text_paragraph(document, TITLE_EN, center=True, size=12, east_asia="Times New Roman", first_line_indent=0)
    add_text_paragraph(document, build_abstract_en(metrics), first_line_indent=0)
    add_text_paragraph(document, "Key Words: " + "; ".join(KEYWORDS_EN), first_line_indent=0)

    add_heading_1(document, "目录")
    for item in [
        "摘要",
        "Abstract",
        "1 绪论",
        "2 数据来源与研究方法",
        "3 全球石油领域研究机构学术影响力评价结果",
        "4 研究主题演化与机构画像分析",
        "5 讨论与建议",
        "6 结论",
        "参考文献",
        "致谢",
    ]:
        add_text_paragraph(document, item, first_line_indent=0)

    document.add_page_break()
    add_heading_1(document, "1 绪论")
    add_heading_2(document, "1.1 研究背景与意义")
    add_text_paragraph(document, "石油仍然是全球能源结构与工业原料体系中的关键支柱。随着非常规油气开发、储层精细表征、二氧化碳利用与封存、催化转化等议题持续升温，科研机构在石油领域中的学术地位、知识生产方式和研究主题布局，已经成为观察行业科技竞争格局的重要窗口。与通用性的大学排名不同，石油领域机构评价不仅涉及论文数量和被引表现，还受到合作网络、工程应用导向以及垂直领域主题契合度的共同影响。")
    add_text_paragraph(document, "从现实需求看，面向全球石油科技竞争与能源安全治理，单纯依靠经验判断已难以支撑机构竞争力识别和资源配置决策。以高质量期刊文献为基础，综合评价全球石油领域研究机构的学术影响力，并进一步追踪热点主题的演化轨迹，有助于识别头部机构优势来源、发现特色机构突破方向，也有助于为高校、企业研发主体和管理部门提供差异化决策参考。")
    add_heading_2(document, "1.2 国内外研究现状")
    add_text_paragraph(document, "现有研究大致可分为两条主线：一条聚焦科研机构或学者学术影响力评价，强调多指标综合评价模型的构建；另一条聚焦主题识别与主题演化分析，强调从大规模文本中抽取研究热点及其变化路径。Schlögl等指出，机构评价结果对数据源选择、计数方式和指标属性设置高度敏感，评价框架的透明性与可解释性是保证结果可信的前提[1]。在国内，林子婕、唐星龙等从多维互动视角讨论学术影响力评价，张璜、高睿则从主题性视角引入LDA模型分析学术影响力问题[3-4]。")
    add_text_paragraph(document, "在石油及相关领域研究中，已有成果更多集中于期刊影响力、单一高校或单一机构群体的科学计量分析。这些研究为垂直领域评价提供了经验，但仍存在全球比较不足、影响力评价与主题演化割裂、机构标准化和跨库清洗规则披露不充分等问题。")
    add_text_paragraph(document, "主题演化研究方面，近年来文献计量与文本挖掘方法结合愈发紧密。王宏宇等基于语义网络识别科研主题演化路径[5]，Blei等提出的LDA模型则为主题发现提供了经典技术框架[12]。然而，面向全球石油领域研究机构的主题演化分析，若缺乏稳定的机构标准化和类型划分，就难以进一步回答“何种机构在何种主题上形成优势、其优势又如何反馈到综合影响力”这一关键问题。")
    add_heading_2(document, "1.3 研究问题与研究目标")
    add_text_paragraph(document, "围绕上述不足，本文重点回答三个问题：第一，如何在高质量期刊文献约束下构建覆盖2011-2025年的全球石油领域正式分析语料，并形成可在论文中合理表述的数据清洗与跨库去重规则；第二，如何建立兼顾科研产出、学术影响和合作国际化的机构综合评价模型，并识别全球石油领域核心研究机构的层级结构；第三，如何将机构评价结果与主题演化结果联系起来，解释不同机构类型和机构层级的主题偏好及其变化。")
    add_text_paragraph(document, "据此，本文的研究目标包括：构建双主键正式语料并完成机构标准化；形成可复现的TOPSIS机构影响力评价结果与层级划分；在三阶段LDA框架下识别石油领域主题演化路径；比较不同机构类型与机构层级的主题差异，并在此基础上提出具有现实针对性的管理建议。")
    add_heading_2(document, "1.4 研究思路与可能创新")
    add_text_paragraph(document, "本文将“数据治理—机构评价—主题演化—策略建议”串联为一条完整分析链条。首先，在原始多源文献基础上执行源内可用性清洗和跨库双主键去重；其次，通过机构译名纠偏、母机构归并和子单元过滤构建标准化机构主表；再次，在Top100核心机构样本上实施组合赋权与TOPSIS评价，并建立机构类型和机构层级画像；最后，复用正式语料对应的主题分析结果，从时间、类型和层级三个维度解释石油领域研究主题的演化逻辑。")
    add_text_paragraph(document, "本文的主要创新点体现在三个方面：一是在论文初稿中显式给出数据清洗、去重与机构标准化规则，增强结果可复查性；二是将机构影响力评价与主题演化分析放在同一正式语料和统一机构口径下讨论；三是根据实际数据可得性对开题报告中的指标体系和类型比较边界进行了收敛，使论文结论更具科学性和可辩护性。")

    document.add_page_break()
    add_heading_1(document, "2 数据来源与研究方法")
    add_heading_2(document, "2.1 数据来源与研究范围")
    add_text_paragraph(document, f"本文将当前持有的三库原始数据统一视为来源于已筛定的109本高质量期刊石油领域文献数据，并以此作为研究的基础数据源。正式分析语料来自CNKI、WOS与CSCD三库2011-2025年记录，经统一清洗和双主键跨库去重后形成正式主表，共包含{metrics['total_docs']}篇文献。按主题建模统计口径，三个阶段的文献数分别为{data['preprocess_rows'][0]['阶段文献数']}篇、{data['preprocess_rows'][1]['阶段文献数']}篇和{data['preprocess_rows'][2]['阶段文献数']}篇。由于本文研究对象是研究机构而非期刊本身，高质量期刊属性在研究中主要承担数据来源约束作用，而不作为后续评价指标。")
    add_heading_2(document, "2.2 数据清洗与跨库去重规则")
    add_text_paragraph(document, "为保证语料构建过程可重复、可审查且可在论文中清晰表述，本文将数据清洗与跨库去重划分为“源内可用性清洗”与“跨库文献标识合并”两个阶段。在源内清洗阶段，CNKI记录不再以DOI是否存在作为入库门槛，而是要求题名、发表时间、期刊和机构四类核心书目信息至少具备可用字段，以避免因中文数据库DOI缺失而无必要地剔除仍具识别价值的文献。")
    add_text_paragraph(document, "在跨库去重阶段，本文采用“双主键”合并策略。第一主键为严格标准DOI键，第二主键为由规范化题名、年份、第一作者和期刊共同构成的元数据回退键。仅当元数据键唯一对应既有分组且不存在DOI冲突时，才允许执行回退合并；弱相似记录、冲突记录和不稳定DOI记录不自动吞并，而是保留在审查表中。")
    add_text_paragraph(document, "完成分组合并后，本文采用来源感知的字段择优融合策略生成正式主表，并显式保留来源库组合、去重依据、标准DOI键和元数据去重键等字段，以支撑后续结果追踪和方法复核。")
    add_heading_2(document, "2.3 机构标准化与研究样本构建")
    add_text_paragraph(document, f"机构标准化遵循“译名可读、标准名统一、颗粒度一致”的原则。对实验室、学院、研究中心、分院等子单元，若能够稳定回溯到母机构，则统一归并到母机构层级；对具有独立学术身份且无法合理归并的机构，则保留独立标准名。经清洗后，机构画像表共保留7882个有效标准化机构，其中企业研发中心{metrics['profile_type_counts'].get('企业研发中心', 0)}个、高校/科研院所{metrics['profile_type_counts'].get('高校/科研院所', 0)}个、政府机构{metrics['profile_type_counts'].get('政府机构', 0)}个、其他机构{metrics['profile_type_counts'].get('其他', 0)}个。")
    add_text_paragraph(document, f"在Top100核心评价样本中，高校/科研院所{metrics['top100_type_counts'].get('高校/科研院所', 0)}家、企业研发中心{metrics['top100_type_counts'].get('企业研发中心', 0)}家、政府机构{metrics['top100_type_counts'].get('政府机构', 0)}家；按层级划分则包括头部引领型{metrics['top100_level_counts'].get('头部引领型', 0)}家、中坚创新型{metrics['top100_level_counts'].get('中坚创新型', 0)}家和特色细分型{metrics['top100_level_counts'].get('特色细分型', 0)}家。")
    add_table_title(document, "表2-1 机构类型分布情况")
    add_table_docx(
        document,
        ["机构类型", "机构数量", "说明"],
        [
            ["企业研发中心", str(metrics["profile_type_counts"].get("企业研发中心", 0)), "标准化后数量最多，反映产业主体广泛参与"],
            ["高校/科研院所", str(metrics["profile_type_counts"].get("高校/科研院所", 0)), "在高影响样本中占据主导"],
            ["政府机构", str(metrics["profile_type_counts"].get("政府机构", 0)), "样本较少，仅作补充观察"],
            ["其他", str(metrics["profile_type_counts"].get("其他", 0)), "未纳入正式类型比较"],
        ],
    )
    add_heading_2(document, "2.4 指标体系与组合赋权方法")
    add_text_paragraph(document, "结合开题报告中“科研产出—学术影响—国际合作”三维框架与当前数据可得性，本文最终保留7项指标进入综合评价模型，即去重论文总数、近五年发文占比、H指数、高被引论文占比、合作论文占比、国际合作论文占比和合作国家/地区数。由于当前重建过程未恢复出可追溯的专家评分数据，本文不再将德尔菲法作为实际赋权手段，而是采用“维度均衡先验权重+熵权法+CRITIC法”的组合赋权方案。")
    add_table_title(document, "表2-2 机构评价指标组合权重")
    add_table_docx(
        document,
        ["指标名称", "所属维度", "组合权重"],
        [[row["label_zh"], row["dimension_zh"], f"{as_float(row['combined_weight']):.6f}"] for row in data["weights"]],
    )
    add_heading_2(document, "2.5 TOPSIS评价与机构分层")
    add_text_paragraph(document, "在综合评价阶段，本文先对全部正向指标进行标准化处理，再依据组合权重构建加权标准化决策矩阵；随后确定正理想解与负理想解，计算各机构到两类理想解的距离，并据此获得TOPSIS综合贴近度得分。得分越高，表示机构越接近理想最优状态。基于综合得分分布，进一步将Top100机构划分为头部引领型、中坚创新型和特色细分型三个层级，用于后续机构画像和主题偏好比较。")
    add_heading_2(document, "2.6 主题演化分析流程")
    add_text_paragraph(document, "主题分析以正式语料中的摘要文本为对象，按2011-2015、2016-2020、2021-2025三个阶段分别建模。文本预处理主要包括摘要筛选、分词、去停用词、石油领域术语归并与词汇表截断。主题数选择采用困惑度与一致性得分联合判定，三个阶段最终均选择11个主题。随后，本文依据相邻阶段主题词分布相似度构建演化连接，以识别主题的延续、分化与迁移关系。")
    add_table_title(document, "表2-3 三阶段主题建模与选模结果")
    add_table_docx(
        document,
        ["阶段", "阶段文献数", "通过领域过滤文献数", "进入建模文献数", "选定主题数", "困惑度", "一致性得分"],
        [
            [
                row["阶段"],
                row["阶段文献数"],
                row["通过领域过滤文献数"],
                row["进入建模文献数"],
                str(metrics["selected_k"][row["阶段"]]),
                metrics["selected_models"][row["阶段"]]["困惑度"],
                metrics["selected_models"][row["阶段"]]["一致性得分"],
            ]
            for row in data["preprocess_rows"]
        ],
    )

    document.add_page_break()
    add_heading_1(document, "3 全球石油领域研究机构学术影响力评价结果")
    add_heading_2(document, "3.1 样本概况与综合格局")
    add_text_paragraph(document, f"依据Top100核心机构综合评价结果，2011-2025年全球石油领域研究机构影响力排名前十位依次为{join_cn(metrics['top10_names'])}。综合排名呈现出明显的“头部集中+多元并存”格局：中国科学院和中国石油大学（北京）构成第一梯队，其后依次为两大石油央企主体以及行业特色高校与综合性大学。")
    add_text_paragraph(document, f"综合排名前20位中除国内头部高校与科研机构外，还出现了{join_cn(metrics['foreign_top20'])}等境外高水平大学，说明国际合作广度和学术影响质量能够帮助部分境外机构在总体发文规模不占绝对优势的情况下进入前列。")
    add_table_title(document, "表3-1 TOPSIS综合排名前十机构")
    add_table_docx(
        document,
        ["排名", "机构名称", "国家/地区", "综合得分", "科研产出子得分", "学术影响子得分", "合作与国际化子得分"],
        [
            [
                row["TOPSIS综合排名"],
                row["标准化机构名称"],
                row["主属国家/地区"],
                row["TOPSIS综合得分"],
                row["科研产出子得分"],
                row["学术影响子得分"],
                row["合作与国际化子得分"],
            ]
            for row in metrics["top10"]
        ],
    )
    add_heading_2(document, "3.2 TOPSIS综合结果分析")
    add_text_paragraph(document, f"从综合得分看，中国科学院以{metrics['top10'][0]['TOPSIS综合得分']}位列第一，其优势来源于高水平的科研产出和稳健的学术影响表现；中国石油大学（北京）以{metrics['top10'][1]['TOPSIS综合得分']}位列第二，在发文规模与持续活跃度上保持显著优势。排名第三和第四的中国石油化工股份有限公司、中国石油天然气股份有限公司表明，大型行业企业已经形成持续稳定的高质量科研输出能力。")
    add_text_paragraph(document, "进一步看，头部机构的优势并不完全相同。中国科学院表现为多维均衡型优势；中国石油大学（北京）体现出主题契合度与长期积累形成的产出优势；中国石油化工股份有限公司和中国石油天然气股份有限公司则显示出企业研发主体在行业问题导向研究上的持续发力。")
    add_heading_2(document, "3.3 分维度优势比较")
    add_text_paragraph(document, f"从科研产出维度看，前五位分别为{join_cn([row['标准化机构名称'] for row in metrics['dimension_grouped']['科研产出维度'][:5]])}；从学术影响维度看，前五位分别为{join_cn([row['标准化机构名称'] for row in metrics['dimension_grouped']['学术影响维度'][:5]])}；从合作与国际化维度看，前五位分别为{join_cn([row['标准化机构名称'] for row in metrics['dimension_grouped']['合作与国际化维度'][:5]])}。")
    add_text_paragraph(document, "这一结果说明，机构综合影响力并不完全等同于单一维度优势。科研产出维度主要受发文规模驱动，国内头部科研院所和行业高校占据主导；学术影响维度中，高质量成果密度较高的高校和国际化高校表现更突出；合作与国际化维度中，境外机构优势更明显，反映出国际合作网络对综合评价结果的支撑作用。")
    add_heading_2(document, "3.4 机构类型与机构层级特征")
    add_text_paragraph(document, "从全部标准化机构画像看，企业研发中心在数量上占绝对多数，但在Top100高影响样本中，高校/科研院所占比更高，说明在高质量期刊口径下，知识生产质量与学术扩散能力仍主要由高校与科研机构主导。与此同时，少数石油央企和大型企业研发主体能够凭借行业问题驱动的稳定产出进入头部样本，体现出鲜明的产业技术牵引特征。")
    add_text_paragraph(document, "按机构层级观察，头部引领型机构在综合得分、合作网络和成果质量等方面表现更为均衡；中坚创新型机构往往在某一维度形成差异化突破；特色细分型机构则更多依赖细分主题长期积累进入核心样本。")

    document.add_page_break()
    add_heading_1(document, "4 研究主题演化与机构画像分析")
    add_heading_2(document, "4.1 三阶段建模结果与热点主题")
    add_text_paragraph(document, f"三阶段LDA建模结果显示，各阶段最优主题数均为11。2011-2015年主题强度最高的三个主题分别为{join_cn([row['主题标签'] for row in metrics['strength_top3']['2011-2015']])}；2016-2020年对应为{join_cn([row['主题标签'] for row in metrics['strength_top3']['2016-2020']])}；2021-2025年对应为{join_cn([row['主题标签'] for row in metrics['strength_top3']['2021-2025']])}。")
    add_text_paragraph(document, "从整体趋势看，石油领域研究热点经历了由传统盆地构造、油田开发和催化转化问题向储层预测、低碳利用与高精度催化位点表征逐步演化的过程。研究重心并非简单替换，而是在传统资源开发议题基础上不断叠加储层精细表征、绿色转型和材料微观机制等新方向。")
    add_figure(document, TOPIC_DIR / "topic_intensity_evolution_curve_dual_key_2011_2025.png", "图4-1 2011-2025年主题强度演化曲线")
    add_heading_2(document, "4.2 主题演化路径分析")
    add_text_paragraph(document, f"在相邻阶段主题相似度匹配基础上，共识别出{len(data['evolution_rows'])}条主题演化路径。较具代表性的路径包括：{metrics['example_paths'][0]}；{metrics['example_paths'][1]}；{metrics['example_paths'][2]}。")
    add_text_paragraph(document, "上述路径表明，石油领域主题演化具有明显的延续性与细化性。一方面，储层表征和开发类研究沿着“宏观构造—储层性质—微观孔隙结构”的路线不断深化；另一方面，催化研究沿着“反应现象—材料活性—位点机理”的方向不断细分。二氧化碳相关主题的持续出现，也表明绿色低碳议题已经成为石油科技研究的重要增长点。")
    add_heading_2(document, "4.3 机构类型主题偏好比较")
    add_text_paragraph(document, f"类型比较结果显示，2011-2015年高校/科研院所占比最高的主题为{join_cn([row['主题标签'] for row in metrics['type_grouped'][('2011-2015', '高校/科研院所')][:2]])}，企业研发中心则更集中于{join_cn([row['主题标签'] for row in metrics['type_grouped'][('2011-2015', '企业研发中心')][:2]])}。2016-2020年，高校/科研院所主要集中于{join_cn([row['主题标签'] for row in metrics['type_grouped'][('2016-2020', '高校/科研院所')][:2]])}，企业研发中心主要集中于{join_cn([row['主题标签'] for row in metrics['type_grouped'][('2016-2020', '企业研发中心')][:2]])}。到2021-2025年，高校/科研院所更偏向{join_cn([row['主题标签'] for row in metrics['type_grouped'][('2021-2025', '高校/科研院所')][:2]])}，企业研发中心则集中于{join_cn([row['主题标签'] for row in metrics['type_grouped'][('2021-2025', '企业研发中心')][:2]])}。")
    add_text_paragraph(document, "由此可见，高校/科研院所更容易在知识探索和前沿主题上保持分布广度，而企业研发中心更偏向与油田开发、储层预测和工程应用直接相关的问题。政府机构样本仅2家，虽然在部分阶段也呈现出与低碳利用、催化转化相关的主题聚集，但由于样本规模过小，本文仅作补充观察而不展开正式比较。")
    add_figure(document, TOPIC_DIR / "topic_distribution_heatmap_by_institution_type_dual_key_2011_2025.png", "图4-2 不同机构类型主题分布热力图")
    add_heading_2(document, "4.4 机构层级主题偏好比较")
    add_text_paragraph(document, f"从机构层级看，2011-2015年头部引领型机构的代表主题为{metrics['level_grouped'][('2011-2015', '头部引领型')][0]['主题标签']}，中坚创新型机构更偏向{metrics['level_grouped'][('2011-2015', '中坚创新型')][0]['主题标签']}；2016-2020年，头部引领型与特色细分型机构都在{metrics['level_grouped'][('2016-2020', '头部引领型')][0]['主题标签']}等储层表征主题上保持活跃，而中坚创新型机构更多集中于{metrics['level_grouped'][('2016-2020', '中坚创新型')][0]['主题标签']}；2021-2025年，头部引领型机构强化了{metrics['level_grouped'][('2021-2025', '头部引领型')][0]['主题标签']}等高复杂度主题，中坚创新型机构则更偏向{metrics['level_grouped'][('2021-2025', '中坚创新型')][0]['主题标签']}。")
    add_text_paragraph(document, "总体而言，机构层级越高，越容易同时覆盖基础研究、工程预测和国际合作关联主题；层级越细分的机构，则越倾向于在少数技术主题上形成高占比优势。这说明影响力层级不仅是结果性排名，也可视为主题布局广度和资源整合能力的侧面反映。")
    add_figure(document, TOPIC_DIR / "topic_distribution_heatmap_by_institution_level_dual_key_2011_2025.png", "图4-3 不同机构层级主题分布热力图")

    document.add_page_break()
    add_heading_1(document, "5 讨论与建议")
    add_heading_2(document, "5.1 机构影响力与主题布局的耦合关系")
    add_text_paragraph(document, "综合前文结果可以发现，石油领域机构影响力并非简单由规模决定，而是由规模、质量、合作和主题布局共同塑造。头部机构往往既拥有较强的科研产出能力，也能覆盖更多前沿主题和国际合作网络；中坚机构更容易在单一维度上形成突破；特色机构则更多通过细分技术路线保持存在感。")
    add_heading_2(document, "5.2 对高校/科研院所的建议")
    add_text_paragraph(document, "高校/科研院所应继续保持在储层精细表征、低碳利用、催化机理等知识探索型主题上的优势，同时加强与企业场景的深度耦合，推动基础研究问题向工程可验证问题延伸。对于已进入头部引领型的机构，应重点提升跨机构协作效率和国际合作深度；对于中坚创新型机构，则应在已有优势主题上形成更清晰的学术品牌。")
    add_heading_2(document, "5.3 对企业研发中心的建议")
    add_text_paragraph(document, "企业研发中心在油田开发、储层预测与工程应用类主题上已形成明显优势，但在高影响成果产出和国际合作方面仍有提升空间。建议企业主体在保持问题导向优势的同时，进一步布局催化材料、二氧化碳利用与封存等前沿议题，强化与高校、科研院所的协同研发机制。")
    add_heading_2(document, "5.4 对管理部门的建议")
    add_text_paragraph(document, "管理部门应推动高质量文献数据、机构标准名和主题监测结果的常态化归集，建立面向重点能源技术方向的滚动监测机制。对于跨机构协同创新，可考虑围绕储层预测、低碳利用和关键催化材料等主题建立联合攻关网络，促进基础研究资源和产业场景资源的有效对接。")
    add_heading_2(document, "5.5 研究局限")
    add_text_paragraph(document, f"本文仍存在三点局限。第一，机构类型识别虽已显著改善，但仍有{metrics['profile_type_counts'].get('其他', 0)}个机构暂列“其他”，说明少量跨国联盟、中心类组织的属性辨识仍需进一步细化。第二，由于缺乏可追溯的专家评分数据，本文未将德尔菲法纳入实际赋权流程，而是采用完全可复现的组合客观赋权方案。第三，经有效机构清洗后，国际组织未形成稳定样本，因此机构类型比较仅正式保留高校/科研院所、企业研发中心和政府机构三类。")

    document.add_page_break()
    add_heading_1(document, "6 结论")
    add_text_paragraph(document, "本文基于109本高质量期刊石油领域文献数据，构建了覆盖2011-2025年的双主键正式语料，并完成了从机构标准化、Top100综合评价到主题演化比较的一体化分析。研究表明，中国科学院、中国石油大学（北京）等机构在综合评价中位居前列，头部格局体现出大型科研院所、行业特色高校与石油企业主体并存的结构特征。")
    add_text_paragraph(document, "主题演化结果显示，石油领域研究热点总体呈现由盆地构造与传统开发问题向储层预测、低碳利用和高精度催化机理演化的趋势。高校/科研院所在知识探索型主题上保持优势，企业研发中心则更集中于工程应用与开发类议题，机构层级越高的样本越容易表现出更宽的主题覆盖度。")
    add_text_paragraph(document, "总体而言，本文为高质量文献驱动下的全球石油领域研究机构评价提供了可复现的技术路线，也为后续机构竞争力诊断、主题监测和资源配置优化提供了数据基础。")

    document.add_page_break()
    add_heading_1(document, "参考文献")
    for index, ref in enumerate(REFERENCES, start=1):
        add_text_paragraph(document, f"[{index}] {ref}", first_line_indent=0)

    add_heading_1(document, "致谢")
    add_text_paragraph(document, "感谢导师在选题、方法与写作过程中的指导，感谢学院提供的研究条件与毕业论文组织支持。本文当前版本为基于最新重建数据链生成的论文初稿，后续仍将结合参考文献细化、正文扩写、图表编号校核和版式套版结果继续完善。")


def write_markdown(markdown_text: str) -> None:
    MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
    MD_OUTPUT.write_text(markdown_text, encoding="utf-8")


def write_docx(markdown_text: str, metrics: dict[str, object], data: dict[str, object]) -> None:
    MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    add_cover_page(document, metrics["cover_info"])
    add_docx_content(document, markdown_text, metrics, data)
    document.save(DOCX_OUTPUT)


def count_zh_chars(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def build_qa(markdown_text: str, metrics: dict[str, object], data: dict[str, object]) -> str:
    refs_count = len(REFERENCES)
    zh_chars = count_zh_chars(markdown_text)
    lines = [
        "# 论文初稿交付检查清单",
        "",
        "生成日期：2026-04-24",
        "",
        "## 已生成文件",
        f"- Markdown 初稿：`{MD_OUTPUT}`",
        f"- Word 初稿：`{DOCX_OUTPUT}`",
        f"- QA 清单：`{QA_OUTPUT}`",
        "",
        "## 自动检查结果",
        f"- 正式语料文献量：`{metrics['total_docs']}`",
        f"- 主题建模文献量：`{metrics['total_model_docs']}`",
        f"- 机构画像表有效机构数：`{len(data['profile_rows'])}`",
        f"- TOPSIS 前十机构表已写入：`是`",
        f"- 组合权重表已写入：`是`",
        f"- 三阶段主题建模表已写入：`是`",
        f"- 图像插入数量：`3`",
        f"- 参考文献条目数：`{refs_count}`",
        f"- Markdown 粗略中文字符数：`{zh_chars}`",
        "",
        "## 当前口径说明",
        "- 正式语料口径为 `2011-2025 dual_key_dedup`。",
        "- 高质量期刊属性作为数据来源约束，不作为机构评价指标本身。",
        "- 机构类型正式比较保留 `高校/科研院所`、`企业研发中心`、`政府机构` 三类。",
        "- `其他` 类型机构未纳入正式类型比较。",
        "- 国际组织未形成稳定样本，已在方法与局限部分显式说明。",
        "",
        "## 仍需人工复核事项",
        "- 参考文献著录格式仍需按学校最终模板逐条核校。",
        "- 目录页码、页眉页脚细化和封面套版建议在最终定稿阶段结合本机 Word 再处理。",
        "- 当前为初稿版本，若需满足管理类专业约 `1.5万字` 的终稿要求，仍建议在讨论、文献综述和案例化论证部分继续扩写。",
        "- 若学校要求附任务书、声明页、翻译附件等前置材料，需在最终提交包中补齐。",
        "",
    ]
    return "\n".join(lines)


def main() -> None:
    data = load_data()
    metrics = build_metrics(data)
    markdown_text = build_markdown(metrics, data)
    write_markdown(markdown_text)
    write_docx(markdown_text, metrics, data)
    qa_text = build_qa(markdown_text, metrics, data)
    QA_OUTPUT.write_text(qa_text, encoding="utf-8")
    print(MD_OUTPUT)
    print(DOCX_OUTPUT)
    print(QA_OUTPUT)


if __name__ == "__main__":
    main()
