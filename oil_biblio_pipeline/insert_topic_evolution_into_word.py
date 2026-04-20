from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Length, Cm


def set_east_asia_font(run, east_asia_font: str, western_font: str = "Times New Roman") -> None:
    run.font.name = western_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def copy_paragraph_format(target, source) -> None:
    target.style = source.style
    target.alignment = source.alignment
    src_fmt = source.paragraph_format
    dst_fmt = target.paragraph_format
    for attr in [
        "first_line_indent",
        "left_indent",
        "right_indent",
        "space_before",
        "space_after",
        "line_spacing",
        "line_spacing_rule",
        "keep_together",
        "keep_with_next",
        "page_break_before",
        "widow_control",
    ]:
        setattr(dst_fmt, attr, getattr(src_fmt, attr))


def set_paragraph_run_style(paragraph, east_asia_font: str, font_size: Pt | None = None, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_east_asia_font(run, east_asia_font)
        if font_size is not None:
            run.font.size = font_size
        if bold is not None:
            run.font.bold = bold


def insert_text_before(anchor, text: str, sample, east_asia_font: str, font_size: Pt | None = None) -> None:
    paragraph = anchor.insert_paragraph_before(text, style=sample.style)
    copy_paragraph_format(paragraph, sample)
    set_paragraph_run_style(paragraph, east_asia_font, font_size=font_size)


def insert_picture_before(anchor, image_path: Path, caption_text: str, caption_sample, source_sample, width: Length = Cm(15.5)) -> None:
    pic_paragraph = anchor.insert_paragraph_before("", style=caption_sample.style)
    copy_paragraph_format(pic_paragraph, caption_sample)
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_paragraph.paragraph_format.first_line_indent = 0
    run = pic_paragraph.add_run()
    run.add_picture(str(image_path), width=width)

    caption_size = caption_sample.runs[0].font.size if caption_sample.runs else Pt(10.5)
    source_size = source_sample.runs[0].font.size if source_sample.runs else Pt(10.5)
    insert_text_before(anchor, caption_text, caption_sample, east_asia_font="宋体", font_size=caption_size)
    insert_text_before(anchor, "资料来源：本文整理", source_sample, east_asia_font="宋体", font_size=source_size)


def find_paragraph(document: Document, exact_text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def find_paragraph_startswith(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph starting with '{prefix}' not found")


def replace_paragraph_text(paragraph, new_text: str, east_asia_font: str, font_size: Pt | None = None) -> None:
    paragraph.text = new_text
    set_paragraph_run_style(paragraph, east_asia_font, font_size=font_size)


def update_document(
    docx_path: Path,
    topic_curve_path: Path,
    institution_type_heatmap_path: Path,
    institution_level_heatmap_path: Path,
) -> None:
    document = Document(docx_path)

    if any(p.text.strip() == "4.6 研究主题动态演化分析" for p in document.paragraphs):
        raise ValueError("Target document already contains the topic evolution section.")

    heading_sample = find_paragraph(document, "4.5 典型机构比较分析")
    body_sample = find_paragraph_startswith(document, "为进一步识别不同机构影响力形成的差异化路径")
    caption_sample = find_paragraph(document, "图4-3 典型机构多指标雷达图")
    source_sample = find_paragraph(document, "资料来源：本文整理")

    paragraph_figure_ref = find_paragraph_startswith(document, "组合权重分布如图 4-5 所示")
    replace_paragraph_text(
        paragraph_figure_ref,
        paragraph_figure_ref.text.replace("图 4-5", "图 3-1"),
        east_asia_font="宋体",
    )
    figure_caption = find_paragraph(document, "图4-5 组合权重分布图")
    caption_size = figure_caption.runs[0].font.size if figure_caption.runs else Pt(10.5)
    replace_paragraph_text(figure_caption, "图3-1 组合权重分布图", east_asia_font="宋体", font_size=caption_size)

    chapter_summary_heading = find_paragraph(document, "4.6 本章小结")
    chapter_summary_heading.text = "4.8 本章小结"
    set_paragraph_run_style(chapter_summary_heading, east_asia_font="黑体")

    chapter_summary_body = find_paragraph_startswith(document, "总体来看，2011-2025 年核心机构综合影响力呈现出以下几个主要特征")
    replace_paragraph_text(
        chapter_summary_body,
        "总体来看，2011-2025 年核心机构综合影响力及研究主题演化呈现出以下几个主要特征：第一，国内科研院所和高水平大学仍然是本领域影响力格局中的主体力量，其中中国科学院和中国石油大学（北京）形成了明显领先的第一梯队；第二，综合影响力的形成机制已不再单纯依赖发文规模，质量表现和合作网络正成为区分机构竞争力的重要变量；第三，境外高水平大学虽然数量不多，但在学术影响和国际合作两个维度上具有显著优势；第四，国内部分行业特色高校已经表现出较强的主题竞争力和合作能力；第五，本领域研究主题正由传统开发生产议题持续转向流动机理、催化转化、非常规储层表征等更加精细化和交叉化的方向；第六，不同机构类型与层级围绕热点主题形成了较为清晰的知识分工，高校/科研院所更偏向基础研究，企业研发中心更偏向工程应用，头部机构则在热点转移早期表现出更强的先发布局能力。",
        east_asia_font="宋体",
    )

    topic_section_body = [
        "为弥补综合影响力评价主要从机构竞争格局出发而对研究内容动态关注不足，本文进一步以正式语料中的论文摘要为基础开展主题演化分析。按照开题报告的技术路线，将样本划分为 2011-2015、2016-2020 和 2021-2025 三个阶段，依次进行分词、去停用词、短语合并、词袋构建和石油领域相关性过滤，并分别训练 LDA 主题模型。各阶段最终进入建模的文献量分别为 74477、49457 和 69293 篇，结合困惑度与主题一致性确定最优主题数分别为 9、10 和 10 个，从而保证了不同阶段主题结构的可比性。",
        "从主题强度变化看，2011-2015 年以“开发/生产/管道/施工”为代表的工程建设与生产运行主题占比最高，2016-2020 年“flow/pressure/gas/parameters”成为最强主题，2021-2025 年则进一步转向“flow/gas/liquid/temperature”和“gas/pressure/pore/water”等围绕流动机理、储层孔隙和非常规油气表征的主题。如图 4-5 所示，主题强度的时序变化并非简单替代，而是呈现出传统开发生产主题逐步弱化、流动表征与储层精细化研究持续增强的演化轨迹。",
        "结合跨阶段主题匹配结果可以进一步发现，催化反应主题沿着“catalyst/reaction/acid/adsorption→catalyst/catalytic→catalyst/catalysts/reaction”的路径持续演进，储层压力与流动主题沿着“pressure/reservoir/water/flow→flow/pressure/gas/parameters→gas/pressure/pore/water”不断深化，储层地质主题则由“gas/reservoir/formation/shale”演化到“reservoir/formation/basin/seismic”。与此同时，2021-2025 年出现了以 hydrate、energy、density、capacity 为核心的新生主题，说明石油领域研究已与非常规资源开发、储能利用和能源转型议题产生更紧密的耦合。",
    ]

    institution_profile_body = [
        "进一步从机构类型看，高校/科研院所、企业研发中心和国际组织在主题分布上呈现出明显分工。2011-2015 年，高校/科研院所更集中于催化反应与吸附等基础研究主题，企业研发中心更偏向开发、生产、管道与施工等应用导向主题，国际组织则更多聚焦储层压力、流动与资源评价等国际通行议题；2016-2020 年以后，高校主题重心转向吸附分离与膜材料，企业与国际组织则共同向页岩储层、孔隙结构和非常规资源开发主题集中；2021-2025 年，高校仍保持对催化转化主题的持续投入，而企业与国际组织则更多围绕气体流动、储层压力与孔隙表征展开研究。如图 4-6 所示，不同机构类型在热点主题上的功能定位并不相同，高校更承担基础研究与方法创新角色，企业更强调工程应用与开发场景，国际组织则在高国际关注度的前沿议题上保持较强活跃度。",
        "从机构层级看，头部引领型、中坚创新型和特色细分型机构也呈现出阶段性差异。2011-2015 年，头部和特色机构主要围绕储层压力与流动主题展开，中坚机构则在催化反应主题上更为集中；2016-2020 年，各层级共同聚焦吸附去除与表面/膜材料主题，说明该阶段形成了较强的共性研究热点；2021-2025 年，各层级的优势主题又同步转向气体压力、孔隙结构与储层表征，显示领域前沿正在出现跨层级扩散。如图 4-7 所示，头部机构更容易在热点转移早期完成布局，中坚与特色机构则通过聚焦细分主题形成差异化跟进路径。",
    ]

    insert_text_before(chapter_summary_heading, "4.6 研究主题动态演化分析", heading_sample, east_asia_font="黑体")
    for paragraph_text in topic_section_body:
        insert_text_before(chapter_summary_heading, paragraph_text, body_sample, east_asia_font="宋体")
    insert_picture_before(
        chapter_summary_heading,
        topic_curve_path,
        "图4-5 主题强度时序演变曲线",
        caption_sample,
        source_sample,
    )

    insert_text_before(chapter_summary_heading, "4.7 机构类型与层级的主题分布差异", heading_sample, east_asia_font="黑体")
    for paragraph_text in institution_profile_body:
        insert_text_before(chapter_summary_heading, paragraph_text, body_sample, east_asia_font="宋体")
    insert_picture_before(
        chapter_summary_heading,
        institution_type_heatmap_path,
        "图4-6 不同机构类型主题分布热力图",
        caption_sample,
        source_sample,
    )
    insert_picture_before(
        chapter_summary_heading,
        institution_level_heatmap_path,
        "图4-7 不同机构层级主题分布热力图",
        caption_sample,
        source_sample,
    )

    conclusion_anchor = find_paragraph(document, "5.2 实践启示")
    insert_text_before(
        conclusion_anchor,
        "第五，研究主题演化表明本领域前沿正由传统开发生产议题向流动机理、催化转化和非常规储层表征持续迁移。三阶段 LDA 结果显示，2011-2015 年以“开发/生产/管道/施工”为代表的工程应用主题占比最高，2016-2020 年转向流动与参数表征主题，2021-2025 年则进一步聚焦流体流动、孔隙结构和储层精细表征，同时出现以 hydrate、energy、density、capacity 为核心的新生主题，说明石油领域研究内容已与非常规资源开发、储能利用和能源转型议题形成更紧密的耦合。",
        body_sample,
        east_asia_font="宋体",
    )
    insert_text_before(
        conclusion_anchor,
        "第六，机构类型与层级之间的主题分工日益清晰。高校/科研院所更偏向催化反应、吸附分离等基础研究主题，企业研发中心更聚焦开发生产与工程应用主题，国际组织则在页岩储层、孔隙结构和高国际关注度议题上更为活跃；从机构层级看，头部机构通常更早完成热点布局，中坚与特色机构则通过聚焦细分方向形成差异化竞争优势。这表明机构影响力不仅体现为规模和质量差异，也体现为在不同研究主题上的功能定位与知识分工。",
        body_sample,
        east_asia_font="宋体",
    )

    future_paragraph = find_paragraph_startswith(document, "未来可以从以下几个方向对本研究进行深化：")
    replace_paragraph_text(
        future_paragraph,
        "未来可以从以下几个方向对本研究进行深化：一是进一步拓展数据来源与样本边界；二是在已完成主题演化分析的基础上，继续强化合作网络、机构-主题耦合网络与更细粒度时间序列比较等结构性和动态性分析；三是扩展评价指标体系；四是推动研究结果与管理实践结合。",
        east_asia_font="宋体",
    )

    chapter5_summary = find_paragraph_startswith(document, "本章在前文数据处理、指标构建和综合评价结果分析的基础上，对 2011-2025 年核心机构综合影响力研究进行了总体归纳。")
    replace_paragraph_text(
        chapter5_summary,
        "本章在前文数据处理、指标构建和综合评价结果分析的基础上，对 2011-2025 年核心机构综合影响力研究进行了总体归纳。研究表明，本研究主题下的机构竞争格局总体呈现国内主导、头部集聚与国际渗透并存的特征，机构影响力形成机制已由单一规模导向逐步转向规模、质量与国际化协同驱动；与此同时，研究主题持续向流动机理、非常规储层表征和能源转型相关议题迁移，不同类型机构则围绕热点主题形成了较为清晰的差异化分工。",
        east_asia_font="宋体",
    )

    document.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Insert topic evolution analysis into an existing thesis Word draft.")
    parser.add_argument("--docx", required=True, help="Target Word file path")
    parser.add_argument("--topic-curve", required=True, help="Topic intensity evolution curve PNG path")
    parser.add_argument("--type-heatmap", required=True, help="Institution type heatmap PNG path")
    parser.add_argument("--level-heatmap", required=True, help="Institution level heatmap PNG path")
    args = parser.parse_args()

    update_document(
        docx_path=Path(args.docx),
        topic_curve_path=Path(args.topic_curve),
        institution_type_heatmap_path=Path(args.type_heatmap),
        institution_level_heatmap_path=Path(args.level_heatmap),
    )
    print(f"updated={args.docx}")


if __name__ == "__main__":
    main()
