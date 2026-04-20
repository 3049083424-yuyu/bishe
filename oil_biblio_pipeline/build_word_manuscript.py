from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_fonts = run._element.rPr.rFonts
    if r_fonts is None:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    normal_element = normal.element
    normal_rpr = normal_element.get_or_add_rPr()
    r_fonts = normal_rpr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        normal_rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "宋体")

    for style_name, size in [("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style_element = style.element
        style_rpr = style_element.get_or_add_rPr()
        style_r_fonts = style_rpr.rFonts
        if style_r_fonts is None:
            style_r_fonts = OxmlElement("w:rFonts")
            style_rpr.append(style_r_fonts)
        style_r_fonts.set(qn("w:eastAsia"), "黑体")

    title = document.styles["Title"]
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(18)


def add_paragraph_with_font(document: Document, text: str, style: str = "Normal", align=None, first_line_indent: bool = True) -> None:
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if style == "Normal":
        paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line_indent else Cm(0)
    run = paragraph.add_run(text)
    set_east_asia_font(run, "宋体" if style == "Normal" else "黑体")


def add_code_block(document: Document, code_lines: list[str]) -> None:
    for line in code_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10.5)
        set_east_asia_font(run, "等线")


def load_figure_map(figure_dir: Path) -> dict[str, tuple[Path, str]]:
    figure_map: dict[str, tuple[Path, str]] = {}
    manifest_path = figure_dir / "figure_manifest_2011_2025.csv"
    if not manifest_path.exists():
        return figure_map

    with manifest_path.open("r", encoding="gb18030", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            figure_no = row["figure_no"].strip()
            file_stem = row["file_stem"].strip()
            title = row["title"].strip()
            png_path = figure_dir / f"{file_stem}.png"
            if png_path.exists():
                figure_map[figure_no] = (png_path, title)
    return figure_map


def add_figure(document: Document, figure_no: str, figure_path: Path, title: str) -> None:
    pic_paragraph = document.add_paragraph()
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_paragraph.add_run()
    run.add_picture(str(figure_path), width=Cm(15.5))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    caption_run = caption.add_run(f"{figure_no} {title}")
    caption_run.font.size = Pt(11)
    set_east_asia_font(caption_run, "宋体")

    source = document.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.paragraph_format.first_line_indent = Cm(0)
    source.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    source_run = source.add_run("资料来源：本文整理")
    source_run.font.size = Pt(10.5)
    set_east_asia_font(source_run, "宋体")


def markdown_to_docx(md_text: str, output_path: Path, figure_dir: Path | None = None) -> None:
    document = Document()
    configure_page(document)
    configure_styles(document)
    figure_map = load_figure_map(figure_dir) if figure_dir else {}
    inserted_figures: set[str] = set()

    lines = md_text.splitlines()
    paragraph_buffer: list[str] = []
    in_code_block = False
    code_lines: list[str] = []
    title_written = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip()).strip()
        paragraph_buffer = []
        if not text:
            return
        if text.startswith("关键词：") or text.startswith("Keywords:"):
            add_paragraph_with_font(document, text, style="Normal", first_line_indent=False)
        else:
            add_paragraph_with_font(document, text, style="Normal")

        for figure_ref in re.findall(r"图\s*4-\d", text):
            figure_no = figure_ref.replace(" ", "")
            if figure_no in figure_map and figure_no not in inserted_figures:
                figure_path, figure_title = figure_map[figure_no]
                add_figure(document, figure_no, figure_path, figure_title)
                inserted_figures.add(figure_no)

    def flush_code_block() -> None:
        nonlocal code_lines
        if code_lines:
            add_code_block(document, code_lines)
            code_lines = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.strip().startswith("```"):
            flush_paragraph()
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_paragraph()
            continue

        if line.startswith("# "):
            flush_paragraph()
            add_paragraph_with_font(document, line[2:].strip(), style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
            title_written = True
            continue
        if line.startswith("## "):
            flush_paragraph()
            if title_written:
                document.add_page_break()
            add_paragraph_with_font(document, line[3:].strip(), style="Heading 1", first_line_indent=False)
            continue
        if line.startswith("### "):
            flush_paragraph()
            add_paragraph_with_font(document, line[4:].strip(), style="Heading 2", first_line_indent=False)
            continue
        if line.startswith("#### "):
            flush_paragraph()
            add_paragraph_with_font(document, line[5:].strip(), style="Heading 3", first_line_indent=False)
            continue

        if line.startswith("- "):
            flush_paragraph()
            add_paragraph_with_font(document, f"• {line[2:].strip()}", style="Normal", first_line_indent=False)
            continue

        paragraph_buffer.append(line)

    flush_paragraph()
    flush_code_block()

    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert a markdown manuscript into a simple Word document.")
    parser.add_argument("--input", required=True, help="Input markdown file path")
    parser.add_argument("--output", required=True, help="Output docx file path")
    parser.add_argument("--figure-dir", help="Optional figure directory for inserting figures")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    figure_dir = Path(args.figure_dir) if args.figure_dir else None

    md_text = input_path.read_text(encoding="utf-8")
    markdown_to_docx(md_text, output_path, figure_dir=figure_dir)
    print(f"output={output_path}")


if __name__ == "__main__":
    main()
